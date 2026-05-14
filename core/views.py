from __future__ import annotations

import copy
import json
import os
import re
import tempfile
import zipfile
from pathlib import Path
from decimal import Decimal, InvalidOperation
from io import BytesIO
import datetime

import urllib.parse
import urllib.request

import openpyxl
import pandas as pd
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.db import IntegrityError, transaction
from django.db.models import Sum, Q, Max
from django.http import FileResponse, HttpResponseForbidden
from django.http import Http404
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_http_methods

from django.utils._os import safe_join

from django.forms import inlineformset_factory
from django.views.decorators.http import require_POST
from .automations.bucket_metrics import analyze_prognosis_workbook, rebuild_projection_with_growth
from .bol_generation import generate_bol_from_form, generate_bol_from_templates
from .forms import (
    BOLForm,
    PricingUploadForm,
    RpcMasterFormatUploadForm,
    TipEntryForm,
    ProjectPlanEntryForm,
    ScheduleActivityForm,
    ScheduleGlobalNoteForm,
    OrderContainerForm,
    OrderContainerLineForm,
    OrderContainerDocumentForm,
    OrderContainerTagForm,
    IndustryRelationshipNodeForm,
    IndustryRelationshipEdgeForm,
)
from .models import (
    Automation,
    Company,
    PricingCustomer,
    PricingQuoteLine,
    TipEntry,
    TipDeposit,
    ProjectPlanEntry,
    ScheduleActivity,
    ScheduleGlobalNote,
    OrderContainer,
    OrderContainerLine,
    OrderContainerDocument,
    OrderContainerTag,
    GardenMap,
    PlantProfile,
    IndustryRelationshipNode,
    IndustryRelationshipEdge,
)

from .rpc_generation import PER_PALLET, generate_rpc_from_form
from .rpcforms import RpcOrderForm
from .services.pricing_import import parse_pricing_matrix_csv
from .services.order_tracker import upsert_container_from_rpc_order
from .context_processors import get_portal_theme
from .services.myshiptracking import bulk_vessel_status
from .services.rpc_master_formatter import parse_rpc_order_xlsx, build_master_format_workbook
from .services.ms_graph_excel import (
    GraphError,
    exchange_code_for_token,
    find_insert_row_for_nld,
    get_access_token_for_user,
    get_authorization_url,
    get_range_values,
    get_used_range,
    insert_range_down,
    parse_excel_date,
    resolve_drive_item_from_share_url,
    set_range_fill,
    set_range_values,
    store_token_for_user,
)
from .services.bucket_color_map import get_bucket_type_argb
import datetime as dt

from dateutil.relativedelta import relativedelta




# -----------------------------
# Pricing helpers
# -----------------------------

def _canon_product_name(raw: str) -> str:
    s = (raw or "").strip().lower()
    if not s:
        return ""
    s = s.replace("+", " plus ")
    s = s.replace("#", " hq ")
    s = s.replace("ltr", "liter")
    s = s.replace("next gen", "nextgen")
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


RPC_PER_PALLET_BY_CANON_NAME = {
    _canon_product_name(name): qty for name, qty in PER_PALLET.items()
}

# Exact pricing-generator labels mapped to the default pieces-per-pallet
# values the user wants to share with the RPC generator workflow.
PRICING_PER_PALLET_OVERRIDES = {
    _canon_product_name("10 Wide Classic HQ#"): 2800,
    _canon_product_name("10 liter classic N6#"): 2800,
    _canon_product_name("10 Wide Classic"): 2800,
    _canon_product_name("10 liter wide Grey"): 2800,
    _canon_product_name("10 liter wide NIR Grey"): 2800,
    _canon_product_name("10 liter wide NextGen"): 2842,
    _canon_product_name("10 liter wide NextGen 2% PrePay discount"): 2842,
    _canon_product_name("5 liter vase"): 6210,
    _canon_product_name("5 liter round"): 3900,
    _canon_product_name("4 liter vase"): 5370,
    _canon_product_name("3 liter round"): 5600,
    _canon_product_name("7 liter vase"): 3240,
    _canon_product_name("8 wide NG"): 4140,
    _canon_product_name("8 wide NG 2% PrePay discount"): 4140,
    _canon_product_name("8 wide Standard"): 3780,
    _canon_product_name("8 liter wide NIR grey"): 3780,
    _canon_product_name("10 Conical"): 3960,
    _canon_product_name("10 Conical 2% PrePay Discount"): 3960,
    _canon_product_name("10 Conical Grey"): 3960,
    _canon_product_name("10 Conical Grey 2% PrePay Discount"): 3960,
    _canon_product_name("10 conical NG"): 4050,
    _canon_product_name("10 Conical NG 2% PrePay discount"): 4050,
    _canon_product_name("10 Conical NG Grey"): 4050,
    _canon_product_name("10 Conical NG Grey 2% PrePay Discount"): 4050,
    _canon_product_name("13 conical"): 2660,
    _canon_product_name("13 Conical Grey"): 2660,
    _canon_product_name("13 Conical NG Grey"): 2730,
    _canon_product_name("13 conical NG"): 2730,
}


def get_rpc_default_pallet_quantity_pieces(product_description: str) -> int:
    """
    Return the default pieces-per-pallet value for a pricing line using the
    same bucket quantity table as the RPC generator, plus explicit pricing
    aliases that should resolve to those same defaults.
    """
    canon = _canon_product_name(product_description)
    if not canon:
        return 0

    override = PRICING_PER_PALLET_OVERRIDES.get(canon)
    if override:
        return int(override)

    exact = RPC_PER_PALLET_BY_CANON_NAME.get(canon)
    if exact:
        return int(exact)

    for known_name, qty in PRICING_PER_PALLET_OVERRIDES.items():
        if known_name and (canon in known_name or known_name in canon):
            return int(qty)

    for known_name, qty in RPC_PER_PALLET_BY_CANON_NAME.items():
        if known_name and (canon in known_name or known_name in canon):
            return int(qty)

    return 0


def normalize_customer_name(raw: str) -> str | None:
    if raw is None:
        return None

    s = str(raw).strip()
    if not s:
        return None

    low = s.lower().strip()

    # Remove the "Los" / LA general pricing customer entirely
    if low in {"los", "los angeles", "los angeles pricing", "la"}:
        return None

    s = re.sub(r"\s+", " ", s).strip()

    # Kendal == Kendal - (remove trailing dash variants)
    s = re.sub(r"\s*-\s*$", "", s).strip()
    low = s.lower()

    mapping = {
        "designers": "Designers Choice",
        "desginers": "Designers Choice",
        "designer's choice": "Designers Choice",
        "designers choice": "Designers Choice",
        "falcon": "Falcon",
        "falcon long": "Falcon",
        "golden": "Golden State",
        "golden state": "Golden State",
        "bandy": "Bandy Ranch",
    }

    if low in mapping:
        return mapping[low]

    return s


def is_euro_customer_name(name: str) -> bool:
    if not name:
        return False

    s = name.strip().lower()
    s = s.replace("'", "")
    s = re.sub(r"\s+", " ", s)

    if s == "pyramid" or s.startswith("pyramid "):
        return True

    if s == "mobis" or s.startswith("mobis "):
        return True
    if s == "mobis flowers" or s.startswith("mobis flowers"):
        return True
    if s.startswith("mobi s"):
        return True

    return False


def get_currency_for_customer_name(name: str):
    if is_euro_customer_name(name):
        return ("EUR", "$")
    return ("USD", "$")


def normalize_destination(customer_name: str, raw_destination: str) -> str:
    if raw_destination is None:
        return ""

    dest = str(raw_destination).strip()
    if not dest:
        return ""

    dest = re.sub(r"\s+", " ", dest).strip()

    # Strip leading "-" artifacts like "- S,C,N" or "-S,C,N"
    dest = re.sub(r"^\-\s*", "", dest).strip()

    cust = (customer_name or "").strip()

    # If destination starts with the customer name, strip it
    if cust:
        cust_re = re.escape(cust)
        dest = re.sub(rf"^{cust_re}\s*-\s*", "", dest, flags=re.IGNORECASE).strip()
        dest = re.sub(rf"^{cust_re}\s+", "", dest, flags=re.IGNORECASE).strip()

    cust_low = cust.lower()

    if cust_low == "designers choice":
        dest = re.sub(r"^choice\s+", "", dest, flags=re.IGNORECASE).strip()

    if cust_low == "golden state":
        dest = re.sub(r"^state\s+", "", dest, flags=re.IGNORECASE).strip()

    if cust_low == "bandy ranch":
        return "Vista"

    # Native fixed destinations
    if cust_low == "native":
        dlow = dest.lower().strip()
        if dlow == "ca":
            return "California"
        if dlow == "co":
            return "Denver"

    dest = re.sub(r"\s+", " ", dest).strip()
    return dest


@login_required
def bucket_projections_zip_export_view(request):
    """
    Downloads BOTH:
      - Bucket_Projections.xlsx
      - Bucket_Adjustments_To_Paste.xlsx
    as a single ZIP.

    IMPORTANT:
      This rebuilds BOTH exports at click-time using the CURRENT session edits:
        - Growth %
        - YoY absolute overrides
        - YoY extra %
        - Customer micro-adjustments

      And it generates a paste-ready list of adjustment line-items (highlighted)
      so you can paste them into the Master List manually.
    """
    tmp_path = request.session.get("bucket_metrics_tmp_path")
    if not tmp_path or not os.path.exists(tmp_path):
        return HttpResponseForbidden("Your uploaded file has expired. Please upload again.")

    # Session edits
    applied_overrides = _get_applied_yoy_overrides(request)
    applied_yoy_pct = _get_applied_yoy_pct_overrides(request)
    applied_customer_deltas = _get_applied_customer_deltas(request)
    applied_growth = _get_applied_growth_overrides(request)

    try:
        # Baseline projection (from the original prognosis workbook, before any tweaks)
        with open(tmp_path, "rb") as fh:
            base_bytes = fh.read()

        baseline_results = analyze_prognosis_workbook(BytesIO(base_bytes))
        baseline_projection_df = baseline_results.get("projection_df")
        if baseline_projection_df is None:
            return HttpResponseForbidden("Could not read baseline projections from the uploaded workbook.")

        # Start from either baseline or the growth-adjusted projection (if growth edits exist)
        if applied_growth:
            projection_df, _yoy_df_unused, _start_label_unused = rebuild_projection_with_growth(
                BytesIO(base_bytes),
                applied_growth,
            )
        else:
            projection_df = baseline_projection_df.copy()

        # Apply all other edits on top
        projection_df = _apply_yoy_overrides_to_projection(projection_df, applied_overrides)
        projection_df = _apply_yoy_pct_overrides_to_projection(projection_df, applied_yoy_pct)
        projection_df = _apply_customer_deltas_to_projection(projection_df, applied_customer_deltas)

        # Always write a FRESH projections export (so the ZIP always reflects current edits)
        projections_path = os.path.join(tempfile.gettempdir(), f"bucket_projections_{request.user.id}.xlsx")
        with pd.ExcelWriter(projections_path, engine="openpyxl") as writer:
            projection_df.to_excel(writer, index=False, sheet_name="Projections")

        request.session["bucket_metrics_projection_export_path"] = projections_path

        # Always write a FRESH adjustments-to-paste export
        adjustments_out = _generate_adjustments_to_paste_workbook(
            tmp_path=tmp_path,
            user_id=request.user.id,
            applied_customer_deltas=applied_customer_deltas,
        )
        if not adjustments_out or not os.path.exists(adjustments_out):
            return HttpResponseForbidden("Could not generate adjustments export.")

        request.session["bucket_metrics_adjustments_export_path"] = adjustments_out
        request.session.modified = True

        adjustments_path = adjustments_out

    except Exception as e:
        return HttpResponseForbidden(f"Could not generate export ZIP: {e}")

    # Build ZIP
    zip_path = os.path.join(tempfile.gettempdir(), f"bucket_exports_{request.user.id}.zip")
    try:
        if os.path.exists(zip_path):
            os.remove(zip_path)
    except Exception:
        pass

    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        z.write(projections_path, arcname="Bucket_Projections.xlsx")
        z.write(adjustments_path, arcname="Bucket_Adjustments_To_Paste.xlsx")

    return FileResponse(
        open(zip_path, "rb"),
        as_attachment=True,
        filename="bucket_exports.zip",
        content_type="application/zip",
    )



def get_or_create_customer_safe(company, name: str):
    """
    Safe under concurrency AND safe inside outer atomic blocks (savepoint).
    """
    try:
        obj = PricingCustomer.objects.get(company=company, name=name)
        return obj, False
    except PricingCustomer.DoesNotExist:
        try:
            with transaction.atomic():
                obj = PricingCustomer.objects.create(company=company, name=name)
            return obj, True
        except IntegrityError:
            obj = PricingCustomer.objects.get(company=company, name=name)
            return obj, False


def upsert_pricing_line_safe(*, company, customer, destination, product_description, price_delivered):
    """
    Safe upsert for PricingQuoteLine unique constraint:
    (company, customer, destination, product_description)
    Uses inner savepoint so IntegrityError doesn't poison outer tx.
    """
    destination = (destination or "").strip()
    product_description = (product_description or "").strip()

    try:
        obj = PricingQuoteLine.objects.get(
            company=company,
            customer=customer,
            destination=destination,
            product_description=product_description,
        )
        created = False
    except PricingQuoteLine.DoesNotExist:
        try:
            with transaction.atomic():
                obj = PricingQuoteLine.objects.create(
                    company=company,
                    customer=customer,
                    destination=destination,
                    product_description=product_description,
                    price_delivered=price_delivered,
                )
            return obj, True
        except IntegrityError:
            obj = PricingQuoteLine.objects.get(
                company=company,
                customer=customer,
                destination=destination,
                product_description=product_description,
            )
            created = False

    if obj.price_delivered != price_delivered:
        obj.price_delivered = price_delivered
        obj.save(update_fields=["price_delivered"])

    return obj, created


@transaction.atomic
def merge_duplicate_pricing_customers(company):
    customers = list(PricingCustomer.objects.filter(company=company).order_by("id"))

    # --- Falcon legacy cleanup: "Beach" -> "Long Beach" ---
    falcon_customers = PricingCustomer.objects.filter(
        company=company,
        name__iexact="Falcon",
    )

    for falcon in falcon_customers:
        bad_lines = PricingQuoteLine.objects.filter(
            company=company,
            customer=falcon,
            destination__iexact="Beach",
        )

        for line in bad_lines:
            existing = PricingQuoteLine.objects.filter(
                company=company,
                customer=falcon,
                destination="Long Beach",
                product_description=line.product_description,
            ).first()

            if existing:
                if existing.price_delivered != line.price_delivered:
                    existing.price_delivered = line.price_delivered
                    existing.save(update_fields=["price_delivered"])
                line.delete()
            else:
                line.destination = "Long Beach"
                line.save(update_fields=["destination"])

    # --- Native legacy cleanup: "CA" -> "California", "CO" -> "Denver" ---
    native_customers = PricingCustomer.objects.filter(
        company=company,
        name__iexact="Native",
    )

    for cust in native_customers:
        lines = PricingQuoteLine.objects.filter(company=company, customer=cust)

        for line in lines:
            dlow = (line.destination or "").strip().lower()

            fixed_dest = None
            if dlow == "ca":
                fixed_dest = "California"
            elif dlow == "co":
                fixed_dest = "Denver"

            if not fixed_dest:
                continue

            if line.destination == fixed_dest:
                continue

            existing = PricingQuoteLine.objects.filter(
                company=company,
                customer=cust,
                destination=fixed_dest,
                product_description=line.product_description,
            ).first()

            if existing:
                if existing.price_delivered != line.price_delivered:
                    existing.price_delivered = line.price_delivered
                    existing.save(update_fields=["price_delivered"])
                line.delete()
            else:
                line.destination = fixed_dest
                line.save(update_fields=["destination"])

    buckets: dict[str, list[PricingCustomer]] = {}
    for c in customers:
        canon = normalize_customer_name(c.name)
        if canon is None:
            buckets.setdefault("__DELETE__", []).append(c)
        else:
            buckets.setdefault(canon, []).append(c)

    for c in buckets.get("__DELETE__", []):
        PricingQuoteLine.objects.filter(company=company, customer=c).delete()
        c.delete()

    for canon_name, cust_list in buckets.items():
        if canon_name == "__DELETE__":
            continue
        if not cust_list:
            continue

        primary = next((c for c in cust_list if (c.name or "").strip() == canon_name), None)
        if primary is None:
            primary = cust_list[0]
            if primary.name != canon_name:
                primary.name = canon_name
                primary.save(update_fields=["name"])

        # Normalize destinations for existing primary lines with collision safety
        for line in PricingQuoteLine.objects.filter(company=company, customer=primary):
            new_dest = normalize_destination(primary.name, line.destination)
            if not new_dest or new_dest == line.destination:
                continue
            try:
                with transaction.atomic():
                    line.destination = new_dest
                    line.save(update_fields=["destination"])
            except IntegrityError:
                existing = PricingQuoteLine.objects.filter(
                    company=company,
                    customer=primary,
                    destination=new_dest,
                    product_description=line.product_description,
                ).first()
                if existing:
                    if existing.price_delivered != line.price_delivered:
                        existing.price_delivered = line.price_delivered
                        existing.save(update_fields=["price_delivered"])
                line.delete()

        duplicates = [c for c in cust_list if c.id != primary.id]

        for dup in duplicates:
            for line in PricingQuoteLine.objects.filter(company=company, customer=dup):
                norm_dest = normalize_destination(primary.name, line.destination)

                existing = PricingQuoteLine.objects.filter(
                    company=company,
                    customer=primary,
                    destination=norm_dest,
                    product_description=line.product_description,
                ).first()

                if existing:
                    if existing.price_delivered != line.price_delivered:
                        existing.price_delivered = line.price_delivered
                        existing.save(update_fields=["price_delivered"])
                    line.delete()
                else:
                    try:
                        with transaction.atomic():
                            line.customer = primary
                            line.destination = norm_dest
                            line.save(update_fields=["customer", "destination"])
                    except IntegrityError:
                        existing2 = PricingQuoteLine.objects.filter(
                            company=company,
                            customer=primary,
                            destination=norm_dest,
                            product_description=line.product_description,
                        ).first()
                        if existing2:
                            if existing2.price_delivered != line.price_delivered:
                                existing2.price_delivered = line.price_delivered
                                existing2.save(update_fields=["price_delivered"])
                        line.delete()

            dup.delete()


# -----------------------------
# Quote-only description overrides (SESSION)
# -----------------------------

def _quote_desc_session_key(company_id: int, customer_id: int) -> str:
    return f"quote_desc_overrides__c{company_id}__cust{customer_id}"


def get_quote_desc_overrides(request, company_id: int, customer_id: int) -> dict:
    """
    Returns { "<line_id>": "New Description", ... } (all string keys for safety).
    """
    key = _quote_desc_session_key(company_id, customer_id)
    data = request.session.get(key, {})
    if isinstance(data, dict):
        return {str(k): str(v) for k, v in data.items() if v is not None}
    return {}


def set_quote_desc_overrides(request, company_id: int, customer_id: int, overrides: dict):
    key = _quote_desc_session_key(company_id, customer_id)
    request.session[key] = {str(k): str(v) for k, v in overrides.items() if v is not None}
    request.session.modified = True


def clear_quote_desc_overrides(request, company_id: int, customer_id: int):
    key = _quote_desc_session_key(company_id, customer_id)
    if key in request.session:
        del request.session[key]
        request.session.modified = True


# -----------------------------
# Prognosis rebuild (with applied tweaks)
# - Adds NEW highlighted lines to "Clean Data (Use for Metrics)"
# - Customer micro adjustments -> "<Customer> Projection"
# - All other deltas -> "General Projection"
# -----------------------------

def _month_label_to_first_of_month(label: str):
    """
    "Jan-26" -> datetime.date(2026, 1, 1)
    Returns None on failure.
    """
    s = (label or "").strip()
    if not s:
        return None
    # Primary format used by the projection table: "Jan-26"
    try:
        dt = pd.to_datetime(s, format="%b-%y")
        return dt.to_pydatetime().date().replace(day=1)
    except Exception:
        pass

    # Be forgiving: some sheets/edits may store months like "Jan-2026", "2026-01", or full dates
    for fmt in ("%b-%Y", "%Y-%m", "%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y"):
        try:
            dt = pd.to_datetime(s, format=fmt)
            return dt.to_pydatetime().date().replace(day=1)
        except Exception:
            continue

    # Last resort: let pandas guess
    try:
        dt = pd.to_datetime(s)
        return dt.to_pydatetime().date().replace(day=1)
    except Exception:
        return None


def _safe_float(v) -> float:
    try:
        if v is None:
            return 0.0
        if isinstance(v, str) and not v.strip():
            return 0.0
        return float(v)
    except Exception:
        return 0.0


def _normalize_month_str(x) -> str:
    # Always compare months as "Mon-YY" labels
    try:
        if x is None:
            return ""
        s = str(x).strip()
        return s
    except Exception:
        return ""


def _build_month_bucket_df(projection_df: pd.DataFrame) -> pd.DataFrame:
    """
    Input: projection_df with month label in first col, bucket columns thereafter.
    Output: DataFrame indexed by month_label (str), columns are bucket names, floats.
    """
    if projection_df is None or projection_df.empty:
        return pd.DataFrame()

    month_col = projection_df.columns[0]
    df = projection_df.copy()

    df[month_col] = df[month_col].astype(str).map(_normalize_month_str)
    df = df.set_index(month_col, drop=True)

    for c in df.columns:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0.0)

    return df


def _append_projection_rows_to_clean_data(
    *,
    workbook_path: str,
    out_path: str,
    customer_rows: list[dict],
    general_rows: list[dict],
):
    wb = openpyxl.load_workbook(workbook_path, data_only=False)

    sheet_name = "Clean Data (Use for Metrics)"
    if sheet_name not in wb.sheetnames:
        ws = wb.create_sheet(sheet_name)
        # Add default headers
        ws.append(["NLD", "RPC#", "City", "Customer", "Bucket Type", "Quantity", "Delivered"])
    else:
        ws = wb[sheet_name]

    # -----------------------------
    # Robust header detection
    # Many of Spencer's prognosis sheets have a couple of title rows, so
    # the actual headers are not guaranteed to be on row 1.
    #
    # We search the first N rows for the row that best matches the expected
    # field headers.
    # -----------------------------
    expected = {"NLD", "Customer", "Bucket Type", "Quantity"}

    def _normalize_header(v) -> str:
        if v is None:
            return ""
        return str(v).strip()

    header_row = 1
    best_score = -1
    scan_rows = min(ws.max_row or 1, 25)
    scan_cols = min(ws.max_column or 1, 200)

    for r in range(1, scan_rows + 1):
        row_vals = {_normalize_header(ws.cell(r, c).value) for c in range(1, scan_cols + 1)}
        score = len({h for h in row_vals if h in expected})
        if score > best_score:
            best_score = score
            header_row = r
        # Early exit if we found all expected headers
        if score >= len(expected):
            header_row = r
            break

    # Map header names -> column index (1-based)
    headers = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(header_row, c).value
        key = _normalize_header(v)
        if key:
            headers[key] = c

    def col_idx(name: str, fallback: int) -> int:
        return headers.get(name, fallback)

    c_nld = col_idx("NLD", 1)
    c_rpc = col_idx("RPC#", 2)
    c_city = col_idx("City", 3)
    c_cust = col_idx("Customer", 4)
    c_bucket = col_idx("Bucket Type", 5)
    c_qty = col_idx("Quantity", 6)
    c_deliv = col_idx("Delivered", 7)

    # Find the last used row (based on key columns), but only scan BELOW the header row
    def row_has_data(r: int) -> bool:
        for cc in (c_nld, c_cust, c_bucket, c_qty):
            v = ws.cell(r, cc).value
            if v is not None and str(v).strip() != "":
                return True
        return False

    last = ws.max_row
    # Trim trailing empties
    while last > header_row and not row_has_data(last):
        last -= 1
    next_row = max(last + 1, header_row + 1)

    highlight_fill = PatternFill("solid", fgColor="FFF2CC")  # light yellow
    italic_font = Font(italic=True, color="000000")


    def write_row(r: int, *, nld, customer, bucket_type, quantity):
        ws.cell(r, c_nld).value = nld
        # Tag these rows so they are easy to find even if highlighting is missed
        ws.cell(r, c_rpc).value = "PROJECTION"  # tag
        ws.cell(r, c_city).value = ""  # blank
        ws.cell(r, c_cust).value = customer
        ws.cell(r, c_bucket).value = bucket_type
        ws.cell(r, c_qty).value = quantity
        ws.cell(r, c_deliv).value = "ADDED"  # tag

        # Highlight + italicize the “projection” line
        for cc in (c_nld, c_rpc, c_city, c_cust, c_bucket, c_qty, c_deliv):
            cell = ws.cell(r, cc)
            cell.fill = highlight_fill
            cell.font = italic_font

    # Write customer micro rows first (so they’re clearly separated)
    for rec in customer_rows:
        nld = rec.get("nld")
        customer = rec.get("customer")
        bucket_type = rec.get("bucket_type")
        qty = rec.get("quantity")

        if nld is None or not customer or not bucket_type:
            continue
        if abs(_safe_float(qty)) < 1e-9:
            continue

        write_row(next_row, nld=nld, customer=customer, bucket_type=bucket_type, quantity=float(qty))
        next_row += 1

    # Write general projection rows
    for rec in general_rows:
        nld = rec.get("nld")
        customer = rec.get("customer")
        bucket_type = rec.get("bucket_type")
        qty = rec.get("quantity")

        if nld is None or not customer or not bucket_type:
            continue
        if abs(_safe_float(qty)) < 1e-9:
            continue

        write_row(next_row, nld=nld, customer=customer, bucket_type=bucket_type, quantity=float(qty))
        next_row += 1



    # -----------------------------
    # Also append projection rows to "Master List".
    #
    # Spencer's "Master List" is a month-grouped table where bucket types are
    # separate columns (there is *no* "Bucket Type" column). Earlier versions
    # of this function tried to find a "Bucket Type" header and silently
    # skipped Master List updates.
    #
    # We now:
    #   1) detect the bucket columns from the main header row
    #   2) locate the subtotal row for the target month (it contains SUM formulas)
    #   3) insert one highlighted row per adjustment *inside that month block*
    #      so it visibly appears under (e.g.) Feb 2026.
    #   4) expand the month subtotal formulas to include the newly inserted rows.
    # -----------------------------
    def _append_to_master_list(ws2, rows: list[dict]):
        if ws2 is None or not rows:
            return

        # Header row: in this workbook it's row 3 (where column A is "NLD" and
        # bucket columns start at I), but we still detect it robustly.
        header_row2 = None
        scan_rows2 = min(ws2.max_row or 1, 40)
        scan_cols2 = min(ws2.max_column or 1, 250)

        def _norm(v):
            return str(v).strip() if v is not None else ""

        for r in range(1, scan_rows2 + 1):
            if _norm(ws2.cell(r, 1).value) == "NLD" and _norm(ws2.cell(r, 2).value) in {"Week #", "Week"}:
                header_row2 = r
                break
        if not header_row2:
            return

        # Column map (fixed meta columns)
        col_nld = 1
        col_week = 2
        col_rpc = 5  # RPC#
        col_city = 6
        col_customer = 7

        # Bucket columns are any non-empty header from column 8 onward on the header row,
        # excluding obvious meta headers.
        meta_names = {
            "NLD",
            "Week #",
            "Week",
            "Departure Date",
            "Client PO#",
            "RPC#",
            "City",
            "Customer",
            "CLASSIC HQ",
            "CLASSIC",
            "NextGen HQ",
            "NextGen N2",
            "5-liter round",
            "5-liter Vase",
        }

        bucket_cols: dict[str, int] = {}
        for c in range(1, scan_cols2 + 1):
            name = _norm(ws2.cell(header_row2, c).value)
            if not name:
                continue
            # We include *all* bucket headers; meta_names is only used to avoid
            # accidentally picking up repeated header blocks.
            if c >= 8 and name not in {"RPC#", "City", "Customer"}:
                bucket_cols[name] = c

        if not bucket_cols:
            return

        highlight_fill2 = PatternFill("solid", fgColor="FFF2CC")
        italic_font2 = Font(italic=True, color="000000")

        # Helper: find the month subtotal row for a given month block.
        # We look for the first SUM formula row after the last datetime row of that month.
        def _find_month_subtotal_row(year: int, month: int) -> int | None:
            last_date_row = None
            for r in range(header_row2 + 1, (ws2.max_row or 1) + 1):
                v = ws2.cell(r, col_nld).value
                if isinstance(v, datetime.datetime) and v.year == year and v.month == month:
                    last_date_row = r
            if last_date_row is None:
                return None
            # scan forward for a SUM formula in first bucket col
            first_bucket_col = min(bucket_cols.values())
            for r in range(last_date_row + 1, min(last_date_row + 120, (ws2.max_row or 1)) + 1):
                v = ws2.cell(r, first_bucket_col).value
                if isinstance(v, str) and v.strip().upper().startswith("=SUM("):
                    return r
            return None

        # Expand SUM formulas in the subtotal row to include newly inserted rows.
        sum_re = re.compile(r"^=SUM\(([A-Z]+)(\d+):([A-Z]+)(\d+)\)\s*$", re.IGNORECASE)

        # Write each adjustment into its month block
        for rec in rows:
            nld = rec.get("nld")
            customer = (rec.get("customer") or "").strip()
            bucket_type = (rec.get("bucket_type") or "").strip()
            qty = _safe_float(rec.get("quantity"))

            if nld is None or not customer or not bucket_type or abs(qty) < 1e-9:
                continue

            # normalize date
            if isinstance(nld, datetime.date) and not isinstance(nld, datetime.datetime):
                nld_dt = datetime.datetime(nld.year, nld.month, nld.day)
            else:
                nld_dt = nld

            if not isinstance(nld_dt, datetime.datetime):
                continue

            subtotal_row = _find_month_subtotal_row(nld_dt.year, nld_dt.month)
            if subtotal_row is None:
                # If we can't locate the month block, fall back to appending to the end
                subtotal_row = (ws2.max_row or 1) + 1

            insert_at = subtotal_row

            # If there is a blank spacer row immediately above the subtotal row, insert
            # ABOVE the spacer so the projection sits right under the last line item
            # for the month (and still above the totals).
            try:
                spacer_row = subtotal_row - 1
                if spacer_row > header_row2:
                    cols_to_check = set([col_nld, col_week, col_rpc, col_city, col_customer]) | set(bucket_cols.values())
                    is_blank = True
                    for cc in cols_to_check:
                        v = ws2.cell(spacer_row, cc).value
                        if v is not None and str(v).strip() != "":
                            is_blank = False
                            break
                    if is_blank:
                        insert_at = spacer_row
            except Exception:
                pass

            if insert_at <= (ws2.max_row or 1):
                ws2.insert_rows(insert_at, amount=1)

            # If we inserted at/above the subtotal row, the subtotal row moved down by 1.
            if insert_at <= subtotal_row:
                subtotal_row = subtotal_row + 1

            # Column A: match Spencer's existing "CUSTOMER PROJECTION" style (string label),
            # not a date. Keep the real customer name in the Customer column.
            ws2.cell(insert_at, col_nld).value = f"{customer.upper()} PROJECTION"
            ws2.cell(insert_at, col_week).value = None
            ws2.cell(insert_at, col_rpc).value = None
            ws2.cell(insert_at, col_city).value = None
            ws2.cell(insert_at, col_customer).value = customer.replace(" Projection", "").strip()

            # Bucket quantity
            if bucket_type in bucket_cols:
                ws2.cell(insert_at, bucket_cols[bucket_type]).value = float(qty)

            # Styling: highlight a reasonable span (meta columns + the one bucket col)
            cols_to_style = {col_nld, col_week, col_rpc, col_city, col_customer}
            if bucket_type in bucket_cols:
                cols_to_style.add(bucket_cols[bucket_type])
            for cc in cols_to_style:
                cell = ws2.cell(insert_at, cc)
                cell.fill = highlight_fill2
                cell.font = italic_font2

            # Expand the subtotal formulas to include this new row if we have a proper subtotal row
            # and it contains SUM(colStart:colEnd) patterns.
            if subtotal_row <= (ws2.max_row or 1):
                for col_name, cidx in bucket_cols.items():
                    v = ws2.cell(subtotal_row, cidx).value
                    if not isinstance(v, str):
                        continue
                    m = sum_re.match(v.strip())
                    if not m:
                        continue
                    col_a, start_row, col_b, end_row = m.group(1).upper(), int(m.group(2)), m.group(3).upper(), int(m.group(4))
                    # Keep the same start, but ensure the end reaches just above subtotal row.
                    new_end = subtotal_row - 1
                    if new_end > end_row:
                        ws2.cell(subtotal_row, cidx).value = f"=SUM({col_a}{start_row}:{col_b}{new_end})"

    ws_master = wb["Master List"] if "Master List" in wb.sheetnames else None
    _append_to_master_list(ws_master, customer_rows)
    _append_to_master_list(ws_master, general_rows)

    # Finally, save the adjusted workbook
    wb.save(out_path)


def _generate_adjustments_to_paste_workbook(
    *,
    tmp_path: str,
    user_id: int,
    applied_customer_deltas: dict,
) -> str | None:
    """Create a small workbook containing ONLY the added adjustment line-items.

    This avoids re-generating / mutating the full prognosis. The output is designed
    so the user can paste rows directly into the Master List.

    Requirements (per user):
      - Column A: "<Month>, <Customer>, PROJECTION" (highlighted)
      - Place bucket quantity in the correct bucket column
      - Column AB (HOLD): quantity
      - Column AC (Bucket Type): bucket name
      - Sorted by month
    """
    if not tmp_path or not os.path.exists(tmp_path):
        return None

    # Load the uploaded prognosis to copy the Master List header structure.
    wb_in = openpyxl.load_workbook(tmp_path)
    if "Master List" not in wb_in.sheetnames:
        return None
    ws_in = wb_in["Master List"]

    # Locate the header row that contains "Bucket Type" (this is the row that
    # defines the exact columns the user wants to paste into).
    header_row = None
    bucket_type_col = None
    max_scan_rows = min(ws_in.max_row or 1, 40)
    max_scan_cols = min(ws_in.max_column or 1, 80)
    for r in range(1, max_scan_rows + 1):
        for c in range(1, max_scan_cols + 1):
            if str(ws_in.cell(r, c).value).strip().lower() == "bucket type":
                header_row = r
                bucket_type_col = c
                break
        if header_row:
            break

    if not header_row:
        return None

    # Build a map of bucket column name -> column index from the header row.
    # Bucket cols are the ones between the first known bucket col and "SUB".
    bucket_cols: dict[str, int] = {}
    header_values: dict[int, str] = {}
    for c in range(1, (ws_in.max_column or max_scan_cols) + 1):
        v = ws_in.cell(header_row, c).value
        if v is None:
            continue
        header_values[c] = str(v).strip()

    # Identify important columns by name
    def _find_col(name: str) -> int | None:
        name_low = name.strip().lower()
        for c, v in header_values.items():
            if str(v).strip().lower() == name_low:
                return c
        return None

    col_nld = _find_col("NLD") or 1
    col_customer = _find_col("Customer")
    col_sub = _find_col("SUB")
    col_hold = _find_col("HOLD")
    col_bucket_type = _find_col("Bucket Type")

    # Bucket columns are everything that looks like a bucket name, excluding meta columns.
    # In your sheet, bucket columns are between the first bucket header and the "SUB" header.
    if col_sub:
        for c in range(1, col_sub):
            v = header_values.get(c)
            if not v:
                continue
            v_low = v.lower()
            if v_low in {"nld", "week #", "departure date", "rpc#", "city", "customer"}:
                continue
            if v_low in {"sub", "hold", "bucket type", "due by"}:
                continue
            # Treat as bucket column
            bucket_cols[v] = c

    # Build output workbook
    wb_out = openpyxl.Workbook()
    ws_out = wb_out.active
    ws_out.title = "Adjustment Lines"

    # Copy column widths (best-effort)
    for c in range(1, (ws_in.max_column or max_scan_cols) + 1):
        col_letter = openpyxl.utils.get_column_letter(c)
        try:
            ws_out.column_dimensions[col_letter].width = ws_in.column_dimensions[col_letter].width
        except Exception:
            pass

    # Copy header row values + styles into row 1
    for c in range(1, (ws_in.max_column or max_scan_cols) + 1):
        src = ws_in.cell(header_row, c)
        dst = ws_out.cell(1, c)
        dst.value = src.value
        try:
            dst._style = copy(src._style)
        except Exception:
            pass

    # Highlight style for added adjustments
    highlight_fill = PatternFill(start_color="FFF59D", end_color="FFF59D", fill_type="solid")
    bold_font = Font(bold=True)

    # Parse applied deltas -> rows
    rows: list[dict] = []
    for key, delta in (applied_customer_deltas or {}).items():
        try:
            month_label, bucket_name, customer_name = key.split("||", 2)
        except ValueError:
            continue

        month_label = str(month_label).strip()
        bucket_name = str(bucket_name).strip()
        customer_name = str(customer_name).strip()
        qty = _safe_float(delta)
        if not month_label or not bucket_name or not customer_name:
            continue
        if abs(qty) < 1e-9:
            continue

        rows.append(
            {
                "month": month_label,
                "bucket": bucket_name,
                "customer": customer_name,
                "qty": qty,
                "nld_sort": _month_label_to_first_of_month(month_label) or datetime.date(1900, 1, 1),
            }
        )

    rows.sort(key=lambda r: (r["nld_sort"], r["customer"], r["bucket"]))

    out_row = 2
    for r in rows:
        month_label = r["month"]
        bucket_name = r["bucket"]
        customer_name = r["customer"]
        qty = r["qty"]

        label = f"{month_label}, {customer_name}, PROJECTION"

        ws_out.cell(out_row, col_nld).value = label

        if col_customer:
            ws_out.cell(out_row, col_customer).value = customer_name

        # Put qty in the correct bucket column
        if bucket_name in bucket_cols:
            ws_out.cell(out_row, bucket_cols[bucket_name]).value = int(round(qty))

  

        # HOLD required
        if col_hold:
            ws_out.cell(out_row, col_hold).value = int(round(qty))

        # Bucket Type required
        if col_bucket_type:
            ws_out.cell(out_row, col_bucket_type).value = bucket_name

        # Highlight the row across A..AC (at least through Bucket Type)
        max_highlight_col = col_bucket_type or 29
        for c in range(1, max_highlight_col + 1):
            cell = ws_out.cell(out_row, c)
            if cell.value is None:
                # Keep blanks as blanks – but still highlight so it stands out.
                pass
            cell.fill = highlight_fill
            # Make the label bold so it stands out when pasted
            if c == col_nld:
                cell.font = bold_font

        out_row += 1

    out_path = os.path.join(tempfile.gettempdir(), f"bucket_adjustments_to_paste_{user_id}.xlsx")
    wb_out.save(out_path)
    return out_path

def _generate_adjusted_prognosis_from_current_session(
    *,
    tmp_path: str,
    user_id: int,
    baseline_projection_df: pd.DataFrame,
    adjusted_projection_df: pd.DataFrame,
    applied_customer_deltas: dict,
) -> str | None:
    """
    Creates a NEW prognosis workbook by appending highlighted projection rows
    to "Clean Data (Use for Metrics)".

    Customer micro adjustments:
      - One line per (Month, Bucket Type, Customer) with Quantity = delta
      - Customer value: "<Customer> Projection"

    General adjustments (YoY + YoY% + Growth, etc):
      - For each (Month, Bucket Type): total_delta - customer_delta_total
      - Customer value: "General Projection"
    """
    if not tmp_path or not os.path.exists(tmp_path):
        return None

    base_df = _build_month_bucket_df(baseline_projection_df)
    adj_df = _build_month_bucket_df(adjusted_projection_df)

    if base_df.empty or adj_df.empty:
        return None

    # Align (months, buckets)
    all_months = sorted(set(base_df.index).union(set(adj_df.index)))
    all_buckets = sorted(set(base_df.columns).union(set(adj_df.columns)))

    base_aligned = base_df.reindex(index=all_months, columns=all_buckets).fillna(0.0)
    adj_aligned = adj_df.reindex(index=all_months, columns=all_buckets).fillna(0.0)

    total_delta = adj_aligned - base_aligned  # month x bucket

    # Aggregate customer deltas into month+bucket totals
    cust_totals = pd.DataFrame(0.0, index=all_months, columns=all_buckets)

    customer_rows: list[dict] = []
    for key, delta in (applied_customer_deltas or {}).items():
        try:
            month_label, bucket_type, customer_name = key.split("||", 2)
        except ValueError:
            continue

        month_label = str(month_label).strip()
        bucket_type = str(bucket_type).strip()
        customer_name = str(customer_name).strip()

        if not month_label or not bucket_type or not customer_name:
            continue

        d = _safe_float(delta)
        if abs(d) < 1e-9:
            continue

        # Add to aggregate
        if month_label in cust_totals.index and bucket_type in cust_totals.columns:
            cust_totals.loc[month_label, bucket_type] = cust_totals.loc[month_label, bucket_type] + d

        nld = _month_label_to_first_of_month(month_label)
        if nld is None:
            continue

        customer_rows.append(
            {
                "nld": nld,
                "customer": f"{customer_name} Projection",
                "bucket_type": bucket_type,
                "quantity": d,
            }
        )

    # Remaining delta is "general projection" delta
    general_delta = total_delta - cust_totals

    general_rows: list[dict] = []
    for month_label in general_delta.index:
        nld = _month_label_to_first_of_month(month_label)
        if nld is None:
            continue
        for bucket_type in general_delta.columns:
            d = _safe_float(general_delta.loc[month_label, bucket_type])
            if abs(d) < 1e-9:
                continue
            general_rows.append(
                {
                    "nld": nld,
                    "customer": "General Projection",
                    "bucket_type": bucket_type,
                    "quantity": d,
                }
            )

    # Sort rows nicely for readability: Month then Customer then Bucket
    def _sort_key(rec: dict):
        # month sort by actual date
        nld = rec.get("nld")
        cust = rec.get("customer") or ""
        b = rec.get("bucket_type") or ""
        return (nld or pd.Timestamp("1900-01-01").date(), cust, b)

    customer_rows.sort(key=_sort_key)
    general_rows.sort(key=_sort_key)

    out_path = os.path.join(tempfile.gettempdir(), f"prognosis_with_projections_{user_id}.xlsx")
    _append_projection_rows_to_clean_data(
        workbook_path=tmp_path,
        out_path=out_path,
        customer_rows=customer_rows,
        general_rows=general_rows,
    )
    return out_path


# -----------------------------
# Core portal views
# -----------------------------

@login_required
def dashboard(request):
    user = request.user
    company = None

    if user.is_superuser:
        automations = Automation.objects.select_related("company").all()
    else:
        try:
            company = Company.objects.get(owner=user)
        except Company.DoesNotExist:
            company = None

        if company:
            automations = Automation.objects.select_related("company").filter(company=company)
        else:
            automations = Automation.objects.none()

    context = {
        "automations": automations,
        "company": company,
        "is_admin": user.is_superuser,
        "portal_theme": get_portal_theme(user=user, company=company),
    }
    return render(request, "core/dashboard.html", context)


def custom_logout(request):
    logout(request)
    return redirect("login")


@login_required
def protected_media_view(request, path: str):
    """Serve user-uploaded media files (PDFs, etc.) through Django.

    Render doesn't automatically serve MEDIA_URL for web services, so we expose
    /media/... via a Django view. This also prevents path traversal.
    """
    try:
        full_path = safe_join(str(settings.MEDIA_ROOT), path)
    except Exception:
        raise Http404("File not found")

    if not os.path.exists(full_path) or not os.path.isfile(full_path):
        raise Http404("File not found")

    # Inline render for PDFs/images, download for everything else
    filename = os.path.basename(full_path)
    resp = FileResponse(open(full_path, "rb"))
    if filename.lower().endswith(".pdf"):
        resp["Content-Disposition"] = f'inline; filename="{filename}"'
    return resp



# -----------------------------
# Microsoft Graph OAuth
# -----------------------------

@login_required
def microsoft_connect_view(request):
    """Start Microsoft OAuth so we can write into the user's OneDrive master workbook."""
    try:
        url = get_authorization_url(state=str(request.user.id))
    except Exception as e:
        messages.error(request, f"Microsoft connect is not configured: {e}")
        return redirect("rpc_master_formatter")
    return redirect(url)


@login_required
def microsoft_callback_view(request):
    """OAuth redirect URI for Microsoft Graph."""
    if request.GET.get("error"):
        err = request.GET.get("error_description") or request.GET.get("error")
        messages.error(request, f"Microsoft authorization failed: {err}")
        return redirect("rpc_master_formatter")

    code = request.GET.get("code")
    if not code:
        messages.error(request, "Missing authorization code from Microsoft.")
        return redirect("rpc_master_formatter")

    try:
        result = exchange_code_for_token(code)
        store_token_for_user(request.user, result)
    except Exception as e:
        messages.error(request, f"Could not save Microsoft connection: {e}")
        return redirect("rpc_master_formatter")

    messages.success(request, "Microsoft connected. You can now insert RPC rows into your OneDrive master.")
    return redirect("rpc_master_formatter")

# -----------------------------
# RPC Order -> Master Spreadsheet Formatter
# -----------------------------

@login_required
@require_http_methods(["GET", "POST"])
def rpc_master_formatter_view(request):
    """
    Upload an RPC order spreadsheet (.xlsx) and return an output workbook
    that matches the row layout of your master spreadsheet.

    Output:
      - headers on row 3
      - data starting row 4
      - one row per bucket type (aggregated totals)
    """
    ms_connected = hasattr(request.user, "ms_graph_token")
    ms_connect_url = reverse("microsoft_connect")

    if request.method == "POST":

        form = RpcMasterFormatUploadForm(request.POST, request.FILES)
        if form.is_valid():
            # With MultipleFileField, this is a list of UploadedFile objects.
            upload_files = form.cleaned_data.get("files") or []
            if not upload_files:
                messages.error(request, "Please choose at least one RPC order spreadsheet.")
                return render(
                    request,
                    "core/rpc_master_formatter.html",
                    {"automation_name": "RPC → Master Formatter", "form": form, "ms_connected": ms_connected, "ms_connect_url": ms_connect_url},
                )

            all_rows = []
            first_meta = None
            for f in upload_files:
                file_bytes = f.read()
                meta, rows = parse_rpc_order_xlsx(file_bytes)
                if first_meta is None:
                    first_meta = meta
                if rows:
                    all_rows.extend(rows)

            # Sort all rows by NLD date (then RPC#, then bucket type) so the output is ready to paste.
            all_rows.sort(
                key=lambda r: (
                    r.nld_date or dt.date(9999, 12, 31),
                    r.rpc_number or 0,
                    (r.bucket_type or "").lower(),
                )
            )

            if not all_rows:
                messages.error(
                    request,
                    "Could not find any line items in the uploaded RPC spreadsheet(s). Make sure they're the standard RPC template.",
                )
                return render(
                    request,
                    "core/rpc_master_formatter.html",
                    {"automation_name": "RPC → Master Formatter", "form": form, "meta": first_meta or {}, "rows": all_rows, "ms_connected": ms_connected, "ms_connect_url": ms_connect_url},
                )

            insert_into_master = request.POST.get("insert_into_master") == "1"
            if insert_into_master and len(upload_files) > 1:
                messages.error(request, "Automatic OneDrive insertion currently supports a single RPC upload at a time. Uncheck the box (or upload one RPC) and generate a combined file instead.")
            elif insert_into_master:
                try:
                    access_token = get_access_token_for_user(request.user)
                    share_url = getattr(settings, "RPC_MASTER_ONEDRIVE_SHARE_URL", None)
                    if not share_url:
                        raise GraphError("RPC_MASTER_ONEDRIVE_SHARE_URL is not set on the server")
                    sheet_name = getattr(settings, "RPC_MASTER_SHEET_NAME", None)
                    ref = resolve_drive_item_from_share_url(access_token, share_url)
                    used = get_used_range(access_token, ref, sheet_name)
                    row_count = used.get("rowCount")
                    if not row_count:
                        # fallback parse from address like Sheet1!A1:AE123
                        addr = (used.get("address") or "")
                        m = re.search(r":\D*(\d+)$", addr)
                        row_count = int(m.group(1)) if m else 0
                    start_row = 4
                    last_row = max(int(row_count or 0), start_row - 1)
                    existing_dates = []
                    if last_row >= start_row:
                        col_vals = get_range_values(access_token, ref, sheet_name, f"A{start_row}:A{last_row}")
                        for row in col_vals:
                            v = row[0] if row else None
                            existing_dates.append(parse_excel_date(v))
                    new_nld = all_rows[0].nld_date
                    if not new_nld:
                        raise GraphError("RPC sheet is missing an NLD date")
                    insert_row = find_insert_row_for_nld(existing_dates, new_nld, start_row)
                    # Insert N blank rows (shift down)
                    for _ in range(len(all_rows)):
                        insert_range_down(access_token, ref, sheet_name, f"A{insert_row}:AE{insert_row}")
                    # Write values + formatting
                    for idx, mr in enumerate(all_rows):
                        rno = insert_row + idx
                        values = [None] * 31
                        values[0] = mr.nld_date.isoformat() if mr.nld_date else None
                        values[1] = mr.nld_week
                        values[3] = int(mr.customer_po) if mr.customer_po and str(mr.customer_po).isdigit() else mr.customer_po
                        values[4] = mr.rpc_number
                        values[5] = mr.city
                        values[6] = mr.customer_name
                        values[7] = mr.mix_flag
                        # quantity into the correct bucket column if known
                        from .services.rpc_master_formatter import BUCKETTYPE_TO_COLUMN, get_transit_times
                        qty_col = BUCKETTYPE_TO_COLUMN.get(mr.bucket_type)
                        if qty_col:
                            values[qty_col - 1] = mr.quantity
                        values[25] = mr.quantity
                        values[26] = mr.bucket_type
                        values[27] = mr.due_by.isoformat() if mr.due_by else None
                        values[28] = mr.due_week
                        avg_days, fast_days = get_transit_times(mr.city)
                        values[29] = avg_days
                        values[30] = fast_days
                        set_range_values(access_token, ref, sheet_name, f"A{rno}:AE{rno}", [values])
                        # light green fills for NLD date + RPC number
                        set_range_fill(access_token, ref, sheet_name, f"A{rno}:A{rno}", "FFE2EFDA")
                        set_range_fill(access_token, ref, sheet_name, f"E{rno}:E{rno}", "FFC6E0B4")
                        bt_fill = get_bucket_type_argb(mr.bucket_type)
                        if bt_fill:
                            set_range_fill(access_token, ref, sheet_name, f"AA{rno}:AA{rno}", bt_fill)
                    messages.success(request, f"Inserted {len(all_rows)} row(s) into your OneDrive master workbook.")
                except Exception as e:
                    messages.error(request, f"Could not insert into OneDrive master: {e}")
            out_bytes = build_master_format_workbook(all_rows)

            # return as downloadable .xlsx
            tmp_path = os.path.join(tempfile.gettempdir(), f"rpc_master_format_{request.user.id}.xlsx")
            with open(tmp_path, "wb") as fh:
                fh.write(out_bytes)

            # If multiple RPCs were uploaded, use a neutral filename.
            rpc_number = (first_meta or {}).get("rpc_number") or "rpc"
            return FileResponse(
                open(tmp_path, "rb"),
                as_attachment=True,
                filename=(f"RPC_{rpc_number}_master_format.xlsx" if len(upload_files) == 1 else "RPC_master_format_combined.xlsx"),
            )
        else:
            # Surface validation issues (e.g., no files selected) so the user
            # doesn't see a "nothing happened" refresh.
            if form.errors:
                messages.error(request, "Please fix the upload form and try again.")
    else:
        form = RpcMasterFormatUploadForm()

    return render(
        request,
        "core/rpc_master_formatter.html",
        {"automation_name": "RPC → Master Formatter", "form": form, "ms_connected": ms_connected, "ms_connect_url": ms_connect_url},
    )


# -----------------------------
# Tip Tracker (Family Automations)
# -----------------------------

@login_required
@require_http_methods(["GET", "POST"])
def tip_tracker_view(request):
    """
    Simple tip tracker:
      - default tip date = today (editable)
      - shift start/end (handles crossing midnight in the model helper)
      - total tips + notes
      - expanded analytics:
          * weekday, start hour, job type
          * weekday×start-hour heatmaps (avg + count)
          * shift length buckets
          * end hour + crossed midnight
          * consistency metrics (median/std/risk score)
          * weekly totals + best/worst + daily trend w/ rolling 30D avg tips/hr
    """
    user = request.user

    # Match your portal pattern: a normal user "owns" exactly one company
    if user.is_superuser:
        company_id = request.GET.get("company_id")
        if company_id:
            company = get_object_or_404(Company, id=company_id)
        else:
            company = Company.objects.order_by("id").first()
    else:
        try:
            company = Company.objects.get(owner=user)
        except Company.DoesNotExist:
            company = None

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    # Two different POSTs hit this route:
    #   1) Add a new entry (main form)
    #   2) Inline edit of a prior entry's tips_total
    #   3) "Deposit" tips since last deposit (piggy bank)
    if request.method == "POST" and request.POST.get("action") == "update_tip":
        entry_id = request.POST.get("entry_id")
        tips_raw = (request.POST.get("tips_total") or "").strip()

        entry = get_object_or_404(TipEntry, id=entry_id, company=company, user=user)

        try:
            entry.tips_total = Decimal(tips_raw) if tips_raw != "" else Decimal("0")
        except (InvalidOperation, ValueError):
            messages.error(request, "Invalid tip amount.")
            return redirect("tip_tracker")

        entry.save(update_fields=["tips_total"])
        messages.success(request, "Tip amount updated.")
        return redirect("tip_tracker")

    if request.method == "POST" and request.POST.get("action") == "deposit_tips":
        # Compute "tips since last deposit" using TipDeposit timestamps.
        last_deposit = (
            TipDeposit.objects
            .filter(company=company, user=user)
            .order_by("-deposited_at")
            .first()
        )
        last_deposit_at = last_deposit.deposited_at if last_deposit else None

        since_qs = TipEntry.objects.filter(company=company, user=user)
        if last_deposit_at:
            since_qs = since_qs.filter(created_at__gt=last_deposit_at)

        tips_since = (since_qs.aggregate(total=Sum("tips_total")).get("total")) or Decimal("0")

        if tips_since <= 0:
            messages.info(request, "No tips to deposit yet.")
            return redirect("tip_tracker")

        TipDeposit.objects.create(company=company, user=user, amount=tips_since)
        messages.success(request, f"Deposited ${tips_since:.2f} 🎉")
        return redirect("tip_tracker")

    if request.method == "POST":
        form = TipEntryForm(request.POST)
        if form.is_valid():
            # -----------------------------------------------------------------
            # Guardrail: prevent accidental double-submits.
            #
            # Sometimes browsers/users will submit the same form twice (double
            # click, lag, mobile tap, etc.). That would create duplicate rows.
            #
            # We store a short-lived fingerprint of the most recent successful
            # submission in the session. If we see the exact same payload again
            # within a couple seconds, we treat it as a duplicate and ignore it.
            # -----------------------------------------------------------------
            try:
                import time
                fp = "|".join(
                    [
                        str(form.cleaned_data.get("tip_date")),
                        str(form.cleaned_data.get("shift_start")),
                        str(form.cleaned_data.get("shift_end")),
                        str(form.cleaned_data.get("tips_total")),
                        str(form.cleaned_data.get("job_type")),
                        (form.cleaned_data.get("notes", "") or "").strip(),
                    ]
                )
                last = request.session.get("tip_tracker_last_submit") or {}
                if (
                    last.get("fp") == fp
                    and isinstance(last.get("ts"), (int, float))
                    and (time.time() - float(last.get("ts"))) < 3.0
                ):
                    messages.info(request, "That entry was already saved.")
                    return redirect("tip_tracker")
                request.session["tip_tracker_last_submit"] = {"fp": fp, "ts": time.time()}
            except Exception:
                # Never block saving tips if session storage fails for any reason.
                pass

            TipEntry.objects.create(
                company=company,
                user=user,
                tip_date=form.cleaned_data["tip_date"],
                shift_start=form.cleaned_data["shift_start"],
                shift_end=form.cleaned_data["shift_end"],
                tips_total=form.cleaned_data["tips_total"],
                notes=form.cleaned_data.get("notes", "") or "",
                job_type=form.cleaned_data["job_type"],
            )
            messages.success(request, "Tip entry saved.")
            return redirect("tip_tracker")
    else:
        form = TipEntryForm()

    entries = (
        TipEntry.objects
        .filter(company=company, user=user)
        .order_by("-tip_date", "-created_at")[:60]
    )

    all_time_qs = TipEntry.objects.filter(company=company, user=user)

    # Deposit state (does NOT affect analytics)
    last_deposit_at = (
        TipDeposit.objects
        .filter(company=company, user=user)
        .aggregate(last=Max("deposited_at"))
        .get("last")
    )

    qs_since = all_time_qs
    if last_deposit_at:
        qs_since = qs_since.filter(created_at__gt=last_deposit_at)

    tips_since_last_deposit = (
        qs_since.aggregate(total=Sum("tips_total")).get("total")
    ) or Decimal("0")

    all_time_deposited = (
        TipDeposit.objects
        .filter(company=company, user=user)
        .aggregate(total=Sum("amount"))
        .get("total")
    ) or Decimal("0")

    grand_total_tips = (
        all_time_qs.aggregate(total=Sum("tips_total")).get("total")
    ) or Decimal("0")

    # All-time average tips per hour = (sum tips) / (sum hours)
    # Duration is derived from shift_start/shift_end, so we compute hours in Python.
    grand_total_hours = sum((e.shift_duration_hours() for e in all_time_qs), 0.0)
    avg_tips_per_hour_all_time = (
        float(grand_total_tips) / grand_total_hours
        if grand_total_hours > 0
        else 0.0
    )

    # Analytics outputs
    weekday_table_html = None
    start_hour_table_html = None
    job_type_table_html = None
    weekday_by_job_table_html = None
    end_hour_table_html = None
    duration_bucket_table_html = None
    heatmap_avg_table_html = None
    heatmap_count_table_html = None
    weekly_table_html = None
    trend_table_html = None
    best_week_summary = None
    worst_week_summary = None
    recommended_shifts_table_html = None


    try:
        qs = TipEntry.objects.filter(company=company, user=user)
        if qs.exists():
            # Build a DataFrame of all entries for analytics.
            rows = []
            for e in qs:
               
                rows.append(
                    {
                        "tip_date": e.tip_date,
                        "weekday": e.tip_date.strftime("%A"),
                        "job_type": e.get_job_type_display(),
                        "start_hour": e.shift_start.hour,
                        "end_hour": e.shift_end.hour,
                        "tips_total": float(e.tips_total),
                        "hours": float(e.shift_duration_hours()),
                        "tips_per_hour": float(e.tips_per_hour()),
                    }
                )

            df = pd.DataFrame(rows)

            # ---- Common formatting helpers ----
            weekday_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
            df["weekday"] = pd.Categorical(df["weekday"], categories=weekday_order, ordered=True)

            def add_consistency_columns(g: pd.DataFrame) -> pd.DataFrame:
                """
                Assumes columns:
                  - avg_tips_per_hour
                  - median_tips_per_hour
                  - std_tips_per_hour
                Adds:
                  - risk_score = std / mean
                """
                g = g.copy()
                g["median_tips_per_hour"] = g["median_tips_per_hour"].round(2)
                g["std_tips_per_hour"] = g["std_tips_per_hour"].fillna(0).round(2)
                g["risk_score"] = (
                    g["std_tips_per_hour"]
                    / g["avg_tips_per_hour"].replace({0: pd.NA})
                ).fillna(0).round(2)
                return g

            # -----------------------------
            # 1) By weekday (with consistency metrics)
            # -----------------------------
            wd = (
                df.groupby("weekday", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips=("tips_total", "mean"),
                    avg_hours=("hours", "mean"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
                .sort_values("weekday")
            )
            wd = add_consistency_columns(wd)
            for c in ["total_tips", "avg_tips", "avg_hours", "avg_tips_per_hour"]:
                wd[c] = wd[c].round(2)
            weekday_table_html = wd.to_html(classes="table table-striped table-sm", index=False, border=0)

            # -----------------------------
            # 2) Most profitable start times (with consistency metrics)
            # -----------------------------
            sh = (
                df.groupby("start_hour", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips=("tips_total", "mean"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
                .sort_values("avg_tips_per_hour", ascending=False)
            )
            sh = add_consistency_columns(sh)
            for c in ["total_tips", "avg_tips", "avg_tips_per_hour"]:
                sh[c] = sh[c].round(2)
            start_hour_table_html = sh.to_html(classes="table table-striped table-sm", index=False, border=0)

            # -----------------------------
            # 3) Breakdown by job type (with consistency metrics)
            # -----------------------------
            jt = (
                df.groupby("job_type", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips=("tips_total", "mean"),
                    avg_hours=("hours", "mean"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
                .sort_values("avg_tips_per_hour", ascending=False)
            )
            jt = add_consistency_columns(jt)
            for c in ["total_tips", "avg_tips", "avg_hours", "avg_tips_per_hour"]:
                jt[c] = jt[c].round(2)
            job_type_table_html = jt.to_html(classes="table table-striped table-sm", index=False, border=0)

            # Optional: job type by weekday
            jtw = (
                df.groupby(["weekday", "job_type"], as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
                .sort_values(["weekday", "avg_tips_per_hour"], ascending=[True, False])
            )
            jtw = add_consistency_columns(jtw)
            for c in ["avg_tips_per_hour"]:
                jtw[c] = jtw[c].round(2)
            weekday_by_job_table_html = jtw.to_html(classes="table table-striped table-sm", index=False, border=0)

            # -----------------------------
            # 4) Shift length buckets (sweet spot)
            # -----------------------------
            df["hours_clamped"] = df["hours"].clip(lower=0)
            df["duration_bucket"] = pd.cut(
                df["hours_clamped"],
                bins=[-0.001, 3, 5, 7, 1000],
                labels=["0–3", "3–5", "5–7", "7+"],
            )
            dur = (
                df.groupby("duration_bucket", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips=("tips_total", "mean"),
                    avg_hours=("hours", "mean"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
            )
            dur = add_consistency_columns(dur)
            for c in ["total_tips", "avg_tips", "avg_hours", "avg_tips_per_hour"]:
                dur[c] = dur[c].round(2)
            duration_bucket_table_html = dur.to_html(classes="table table-striped table-sm", index=False, border=0)

            # -----------------------------
            # 5) Start vs end time + crossed midnight
            # -----------------------------
            eh = (
                df.groupby("end_hour", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    median_tips_per_hour=("tips_per_hour", "median"),
                    std_tips_per_hour=("tips_per_hour", "std"),
                )
                .sort_values("avg_tips_per_hour", ascending=False)
            )
            eh = add_consistency_columns(eh)
            for c in ["avg_tips_per_hour"]:
                eh[c] = eh[c].round(2)
            end_hour_table_html = eh.to_html(classes="table table-striped table-sm", index=False, border=0)


            # -----------------------------
            # 6) Heatmap: weekday × start hour (avg tips/hr + count)
            # -----------------------------
            heat_avg = (
                df.pivot_table(
                    index="weekday",
                    columns="start_hour",
                    values="tips_per_hour",
                    aggfunc="mean",
                )
                .reindex(weekday_order)
            )
            heat_cnt = (
                df.pivot_table(
                    index="weekday",
                    columns="start_hour",
                    values="tips_per_hour",
                    aggfunc="count",
                )
                .reindex(weekday_order)
            )
            heat_avg = heat_avg.round(2)

            if len(heat_avg.columns) > 0:
                heat_avg = heat_avg.reindex(sorted(heat_avg.columns), axis=1)
                heat_cnt = heat_cnt.reindex(sorted(heat_cnt.columns), axis=1)

            heatmap_avg_table_html = heat_avg.to_html(classes="table table-striped table-sm", border=0)
            heatmap_count_table_html = heat_cnt.to_html(classes="table table-striped table-sm", border=0)
            # -----------------------------
            # 6b) Recommended shifts (weekday × start hour)
            # -----------------------------
            MIN_SHIFTS_PER_CELL = 3   # bump to 5 later once you have more data
            TOP_N = 10

            # Flatten the pivot tables to rows: weekday, start_hour, avg_tips_per_hour, shifts
            rec = (
                heat_avg.stack(dropna=False)
                .rename("avg_tips_per_hour")
                .reset_index()
                .merge(
                    heat_cnt.stack(dropna=False).rename("shifts").reset_index(),
                    on=["weekday", "start_hour"],
                    how="left",
                )
            )

            # Clean + filter
            rec["shifts"] = rec["shifts"].fillna(0).astype(int)
            rec = rec.dropna(subset=["avg_tips_per_hour"])
            rec = rec[rec["shifts"] >= MIN_SHIFTS_PER_CELL].copy()

            # Sort best first, and format start_hour nicely
            rec = rec.sort_values(["avg_tips_per_hour", "shifts"], ascending=[False, False]).head(TOP_N)
            rec["avg_tips_per_hour"] = rec["avg_tips_per_hour"].round(2)
            rec["start_time"] = rec["start_hour"].astype(int).astype(str).str.zfill(2) + ":00"

            # Display columns
            rec_display = rec[["weekday", "start_time", "shifts", "avg_tips_per_hour"]].rename(
                columns={
                    "weekday": "Weekday",
                    "start_time": "Start",
                    "shifts": "Shifts (sample)",
                    "avg_tips_per_hour": "Avg tips/hr",
                }
            )

            recommended_shifts_table_html = rec_display.to_html(
                classes="table table-striped table-sm",
                index=False,
                border=0,
            )

            # -----------------------------
            # 7) Weekly totals + best/worst week + rolling averages (7-shift, 30-day)
            # -----------------------------
            df2 = df.copy()
            df2["tip_date"] = pd.to_datetime(df2["tip_date"])
            df2 = df2.sort_values(["tip_date", "start_hour"], ascending=[True, True])

            # Weekly summary (week starts Monday)
            df2["week_start"] = df2["tip_date"].dt.to_period("W-MON").apply(lambda p: p.start_time.date())
            wk = (
                df2.groupby("week_start", as_index=False)
                .agg(
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                    avg_hours=("hours", "mean"),
                )
                .sort_values("week_start", ascending=False)
            )
            for c in ["total_tips", "avg_tips_per_hour", "avg_hours"]:
                wk[c] = wk[c].round(2)
            weekly_table_html = wk.to_html(classes="table table-striped table-sm", index=False, border=0)

            if not wk.empty:
                best_row = wk.sort_values(["avg_tips_per_hour", "total_tips"], ascending=[False, False]).iloc[0]
                worst_row = wk.sort_values(["avg_tips_per_hour", "total_tips"], ascending=[True, True]).iloc[0]
                best_week_summary = {
                    "week_start": str(best_row["week_start"]),
                    "shifts": int(best_row["shifts"]),
                    "total_tips": float(best_row["total_tips"]),
                    "avg_tips_per_hour": float(best_row["avg_tips_per_hour"]),
                }
                worst_week_summary = {
                    "week_start": str(worst_row["week_start"]),
                    "shifts": int(worst_row["shifts"]),
                    "total_tips": float(worst_row["total_tips"]),
                    "avg_tips_per_hour": float(worst_row["avg_tips_per_hour"]),
                }

            # Rolling averages (daily)
            daily = (
                df2.groupby(df2["tip_date"].dt.date, as_index=False)
                .agg(
                    date=("tip_date", "first"),
                    shifts=("tips_total", "count"),
                    total_tips=("tips_total", "sum"),
                    avg_tips_per_hour=("tips_per_hour", "mean"),
                )
                .sort_values("date")
            )
            daily["avg_tips_per_hour"] = daily["avg_tips_per_hour"].round(2)
            daily["total_tips"] = daily["total_tips"].round(2)
            daily["rolling_30_day_avg_tips_per_hour"] = (
                daily.set_index("date")["avg_tips_per_hour"]
                .rolling("30D", min_periods=1)
                .mean()
                .round(2)
                .values
            )

            trend = daily.tail(45).copy()
            trend["date"] = trend["date"].dt.date.astype(str)
            trend_table_html = trend.to_html(classes="table table-striped table-sm", index=False, border=0)

    except Exception as e:
        messages.warning(request, f"Analytics error: {e}")

    context = {
        "automation_name": "Tip Tracker",
        "company": company,
        "form": form,
        "entries": entries,
        "grand_total_tips": grand_total_tips,
        "tips_since_last_deposit": tips_since_last_deposit,
        "all_time_deposited": all_time_deposited,
        "last_deposit_at": last_deposit_at,
        "avg_tips_per_hour_all_time": avg_tips_per_hour_all_time,
        "weekday_table_html": weekday_table_html,
        "start_hour_table_html": start_hour_table_html,
        "job_type_table_html": job_type_table_html,
        "weekday_by_job_table_html": weekday_by_job_table_html,
        "end_hour_table_html": end_hour_table_html,
        "duration_bucket_table_html": duration_bucket_table_html,
        "heatmap_avg_table_html": heatmap_avg_table_html,
        "heatmap_count_table_html": heatmap_count_table_html,
        "weekly_table_html": weekly_table_html,
        "trend_table_html": trend_table_html,
        "best_week_summary": best_week_summary,
        "worst_week_summary": worst_week_summary,
        "recommended_shifts_table_html": recommended_shifts_table_html,

    }
    return render(request, "core/tip_tracker.html", context)


@login_required
@require_http_methods(["GET"])
def tip_tracker_export_excel(request):
    """
    Export a polished Excel tip report that also prints cleanly to PDF.
    Includes:
      - Summary (grand total, averages)
      - Recent entries table (all entries for the user/company)
      - Key analytics tables already computed in tip_tracker_view (weekday/job type)
    """
    user = request.user

    company_id = request.GET.get("company_id")
    if company_id:
        company = get_object_or_404(Company, pk=company_id)
    else:
        company = Company.objects.first()

    # Safety: enforce company scoping if your app requires it
    # If you have per-user company access rules, apply them here.
    if company is None:
        return HttpResponseForbidden("No company found.")

    all_time_qs = TipEntry.objects.filter(company=company, user=user).order_by("tip_date", "created_at")

    # Build a DataFrame for nicer grouping/analytics
    rows = []
    for e in all_time_qs:
        duration = e.shift_duration_hours()
        tips_total = float(e.tips_total or 0)
        tips_hr = float(e.tips_per_hour() or 0)
        rows.append({
            "Date": e.tip_date,
            "Job type": e.get_job_type_display() if hasattr(e, "get_job_type_display") else getattr(e, "job_type", ""),
            "Start": e.shift_start,
            "End": e.shift_end,
            "Hours": duration,
            "Tips": tips_total,
            "Tips/hr": tips_hr,
            "Notes": e.notes or "",
        })

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Weekday"] = pd.to_datetime(df["Date"]).dt.day_name()

    grand_total = Decimal("0")
    try:
        grand_total = (all_time_qs.aggregate(total=Sum("tips_total")).get("total")) or Decimal("0")
    except Exception:
        grand_total = Decimal("0")

    avg_tips = float(df["Tips"].mean()) if not df.empty else 0.0
    avg_tips_hr = float(df["Tips/hr"].mean()) if not df.empty else 0.0
    total_shifts = int(len(df)) if not df.empty else 0

    # --- Workbook ---
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tip Report"

    # Print/PDF friendly page setup
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 0.35
    ws.page_margins.right = 0.35
    ws.page_margins.top = 0.5
    ws.page_margins.bottom = 0.5

    thin = Side(style="thin", color="1F2937")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    header_fill = PatternFill("solid", fgColor="F3F4F6")  # light gray
    accent_fill = PatternFill("solid", fgColor="FACC15")

    title_font = Font(size=18, bold=True, color="FACC15")
    h_font = Font(size=11, bold=True, color="000000")
    normal_font = Font(size=11, color="000000")

    muted_font = Font(size=10, color="000000")


    # Background columns (make sheet look consistent with your dark UI)
    # Excel doesn't support whole-sheet background well, so we style headers + key cells.

    # --- Title / Summary ---
    ws["A1"] = "Tip Tracker Report"
    ws["A1"].font = title_font
    ws.merge_cells("A1:H1")

    ws["A2"] = f"Company: {company.name}"
    ws["A2"].font = muted_font
    ws.merge_cells("A2:H2")

    ws["A3"] = f"Generated: {timezone.localtime(timezone.now()).strftime('%Y-%m-%d %I:%M %p')}"
    ws["A3"].font = muted_font
    ws.merge_cells("A3:H3")

    # KPI row
    kpis = [
        ("Grand total tips", float(grand_total)),
        ("Total shifts", total_shifts),
        ("Avg tips/shift", avg_tips),
        ("Avg tips/hr", avg_tips_hr),
    ]
    start_row = 5
    col = 1
    for label, value in kpis:
        c1 = ws.cell(row=start_row, column=col, value=label)
        c2 = ws.cell(row=start_row + 1, column=col, value=value)
        ws.merge_cells(start_row=start_row, start_column=col, end_row=start_row, end_column=col+1)
        ws.merge_cells(start_row=start_row+1, start_column=col, end_row=start_row+1, end_column=col+1)

        c1 = ws.cell(row=start_row, column=col)
        c2 = ws.cell(row=start_row + 1, column=col)

        c1.fill = header_fill
        c1.font = h_font
        c1.alignment = Alignment(horizontal="center", vertical="center")
        c1.border = border

        c2.fill = PatternFill("solid", fgColor="111827")
        c2.font = Font(size=14, bold=True, color="E5E7EB")
        c2.number_format = "#,##0.00" if isinstance(value, float) else "0"
        c2.alignment = Alignment(horizontal="center", vertical="center")
        c2.border = border

        # apply border to merged partner cell
        ws.cell(row=start_row, column=col+1).border = border
        ws.cell(row=start_row+1, column=col+1).border = border

        col += 2

    current_row = start_row + 3

    # --- Entries table ---
    ws["A{}".format(current_row)] = "Entries"
    ws["A{}".format(current_row)].font = Font(size=13, bold=True, color="E5E7EB")
    ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=8)
    current_row += 1

    columns = ["Date", "Weekday", "Job type", "Start", "End", "Hours", "Tips", "Tips/hr", "Notes"]
    # We'll print Notes last but keep width reasonable by truncating.
    columns = ["Date", "Weekday", "Job type", "Start", "End", "Hours", "Tips", "Tips/hr", "Notes"]

    # Header row
    for ci, name in enumerate(columns, start=1):
        cell = ws.cell(row=current_row, column=ci, value=name)
        cell.fill = header_fill
        cell.font = h_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    current_row += 1

    def _safe_note(s: str, max_len: int = 120) -> str:
        s = (s or "").strip()
        if len(s) <= max_len:
            return s
        return s[:max_len-1] + "…"

    # Data rows
    if df.empty:
        ws.cell(row=current_row, column=1, value="No entries yet.").font = muted_font
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
        current_row += 2
    else:
        for _, r in df.iterrows():
            values = [
                r.get("Date"),
                r.get("Weekday"),
                r.get("Job type"),
                r.get("Start"),
                r.get("End"),
                float(r.get("Hours") or 0),
                float(r.get("Tips") or 0),
                float(r.get("Tips/hr") or 0),
                _safe_note(str(r.get("Notes") or "")),
            ]
            for ci, v in enumerate(values, start=1):
                cell = ws.cell(row=current_row, column=ci, value=v)
                cell.font = normal_font
                cell.border = border
                cell.alignment = Alignment(horizontal="left" if ci in (3,9) else "center", vertical="top", wrap_text=(ci==9))
                if ci in (6,7,8):
                    cell.number_format = "#,##0.00"
            current_row += 1
        current_row += 2

    # Freeze panes at the entries header
    ws.freeze_panes = "A{}".format(start_row + 3)

    # Column widths (PDF friendly)
    widths = {
        1: 11,  # Date
        2: 12,  # Weekday
        3: 22,  # Job type
        4: 9,   # Start
        5: 9,   # End
        6: 9,   # Hours
        7: 10,  # Tips
        8: 10,  # Tips/hr
        9: 40,  # Notes
    }
    for c, w in widths.items():
        ws.column_dimensions[get_column_letter(c)].width = w

    # --- Quick analytics (if data exists) ---
    if not df.empty:
        # Weekday summary
        ws["A{}".format(current_row)] = "Weekday summary"
        ws["A{}".format(current_row)].font = Font(size=13, bold=True, color="E5E7EB")
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
        current_row += 1

        wd = (
            df.groupby("Weekday")
              .agg(shifts=("Tips", "count"), total_tips=("Tips", "sum"), avg_tips=("Tips", "mean"), avg_tips_hr=("Tips/hr", "mean"))
              .reset_index()
        )
        wd_cols = ["Weekday", "shifts", "total_tips", "avg_tips", "avg_tips_hr"]
        for ci, name in enumerate(wd_cols, start=1):
            cell = ws.cell(row=current_row, column=ci, value=name.replace("_", " ").title())
            cell.fill = header_fill
            cell.font = h_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        current_row += 1

        for _, r in wd.iterrows():
            vals = [r["Weekday"], int(r["shifts"]), float(r["total_tips"]), float(r["avg_tips"]), float(r["avg_tips_hr"])]
            for ci, v in enumerate(vals, start=1):
                cell = ws.cell(row=current_row, column=ci, value=v)
                cell.font = normal_font
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if ci >= 3:
                    cell.number_format = "#,##0.00"
            current_row += 1

        current_row += 2

        # Job type summary
        ws["A{}".format(current_row)] = "Job type summary"
        ws["A{}".format(current_row)].font = Font(size=13, bold=True, color="E5E7EB")
        ws.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
        current_row += 1

        jt = (
            df.groupby("Job type")
              .agg(shifts=("Tips", "count"), total_tips=("Tips", "sum"), avg_tips=("Tips", "mean"), avg_tips_hr=("Tips/hr", "mean"))
              .reset_index()
              .sort_values("avg_tips_hr", ascending=False)
        )
        jt_cols = ["Job type", "shifts", "total_tips", "avg_tips", "avg_tips_hr"]
        for ci, name in enumerate(jt_cols, start=1):
            cell = ws.cell(row=current_row, column=ci, value=name.replace("_", " ").title())
            cell.fill = header_fill
            cell.font = h_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        current_row += 1

        for _, r in jt.iterrows():
            vals = [r["Job type"], int(r["shifts"]), float(r["total_tips"]), float(r["avg_tips"]), float(r["avg_tips_hr"])]
            for ci, v in enumerate(vals, start=1):
                cell = ws.cell(row=current_row, column=ci, value=v)
                cell.font = normal_font
                cell.border = border
                cell.alignment = Alignment(horizontal="left" if ci == 1 else "center", vertical="center")
                if ci >= 3:
                    cell.number_format = "#,##0.00"
            current_row += 1

        current_row += 1

    # Set print area to used range
    ws.print_area = f"A1:{get_column_letter(ws.max_column)}{ws.max_row}"

    # Repeat header rows (title + company + generated + blank + KPI labels/values)
    ws.print_title_rows = "1:6"

    out = BytesIO()
    wb.save(out)
    out.seek(0)

    filename = f"tip_report_{company.name}_{timezone.localtime(timezone.now()).strftime('%Y%m%d_%H%M')}.xlsx"
    filename = re.sub(r"[^A-Za-z0-9_.-]+", "_", filename)

    return FileResponse(
        out,
        as_attachment=True,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

@login_required
@require_http_methods(["GET", "POST"])
def tip_entry_delete_view(request, entry_id: int):
    # Be forgiving: if the entry doesn't exist (e.g., stale UI / double-submit race),
    # don't throw a 404 page — just return the user to the tracker with a message.
    entry = TipEntry.objects.filter(id=entry_id).first()
    if not entry:
        messages.info(request, "That tip entry no longer exists.")
        return redirect("tip_tracker")

    # Permission check: owner of company (or superuser), and must match user+company
    if not (
        request.user.is_superuser
        or (entry.company and entry.company.owner == request.user)
    ):
        return HttpResponseForbidden("You are not allowed to delete this entry.")

    # Optional: also restrict to same user who created it (recommended for Hailey’s use case)
    if not request.user.is_superuser and entry.user_id != request.user.id:
        return HttpResponseForbidden("You are not allowed to delete this entry.")

    if request.method == "POST":
        confirm = (request.POST.get("confirm") or "").strip().lower()
        if confirm == "yes":
            entry.delete()
            messages.success(request, "Entry deleted.")
            return redirect("tip_tracker")

        messages.info(request, "Deletion cancelled.")
        return redirect("tip_tracker")

    # GET = warning/confirmation page
    return render(request, "core/tip_entry_confirm_delete.html", {"entry": entry})

# -----------------------------
# YoY apply/unapply logic
# - Apply YoY: replace projection cell with prev-year absolute value
# - Optional extra %: applied on top of that replacement value
# - Customer micro-deltas: add/subtract absolute units to a month+bucket
# -----------------------------

def _yoy_session_key(month_label: str, col_name: str) -> str:
    return f"{month_label}||{col_name}"


def _get_applied_yoy_overrides(request) -> dict:
    """
    Stores absolute replacement values:
      { "Dec-25||CLASSIC HQ": 1234.0, ... }
    """
    data = request.session.get("bucket_metrics_applied_yoy", {})
    if not isinstance(data, dict):
        return {}

    cleaned: dict[str, float] = {}
    for k, v in data.items():
        try:
            cleaned[str(k)] = float(v)
        except Exception:
            continue
    return cleaned


def _set_applied_yoy_overrides(request, overrides: dict) -> None:
    request.session["bucket_metrics_applied_yoy"] = {str(k): float(v) for k, v in overrides.items()}
    request.session.modified = True


def _get_applied_yoy_pct_overrides(request) -> dict:
    """
    Stores extra percent to apply AFTER YoY replacement:
      { "Dec-25||CLASSIC HQ": 0.10, ... }  # 10% as decimal
    """
    data = request.session.get("bucket_metrics_applied_yoy_pct", {})
    if not isinstance(data, dict):
        return {}

    cleaned: dict[str, float] = {}
    for k, v in data.items():
        try:
            cleaned[str(k)] = float(v)
        except Exception:
            continue
    return cleaned


def _set_applied_yoy_pct_overrides(request, overrides: dict) -> None:
    request.session["bucket_metrics_applied_yoy_pct"] = {str(k): float(v) for k, v in overrides.items()}
    request.session.modified = True

def _get_applied_growth_overrides(request) -> dict:
    """
    Stores growth % values (as decimals) keyed by the REAL bucket column name:
      { "CLASSIC HQ": 0.05, "10 Conical": -0.02, ... }
    This lets us rebuild the same "apply_growth" projection later (exports, ZIP, etc).
    """
    data = request.session.get("bucket_metrics_applied_growth", {})
    if not isinstance(data, dict):
        return {}

    cleaned: dict[str, float] = {}
    for k, v in data.items():
        try:
            cleaned[str(k)] = float(v)
        except Exception:
            continue
    return cleaned


def _set_applied_growth_overrides(request, overrides: dict) -> None:
    request.session["bucket_metrics_applied_growth"] = {str(k): float(v) for k, v in overrides.items()}
    request.session.modified = True




def _apply_yoy_overrides_to_projection(projection_df, applied_overrides: dict):
    """
    Replace the projection value with the stored absolute value.
    """
    if projection_df is None or projection_df.empty:
        return projection_df

    month_col = projection_df.columns[0]
    df = projection_df.copy()

    for key, target_value in (applied_overrides or {}).items():
        try:
            month_label, col_name = key.split("||", 1)
        except ValueError:
            continue

        if col_name not in df.columns:
            continue

        mask = df[month_col].astype(str) == str(month_label)
        if not mask.any():
            continue

        try:
            df.loc[mask, col_name] = float(target_value)
        except Exception:
            continue

    # round to whole numbers for display
    for c in df.columns[1:]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).round(0).astype(int)

    return df


def _apply_yoy_pct_overrides_to_projection(projection_df, yoy_pct_overrides: dict):
    """
    Applies extra % AFTER YoY replacement:
      value := value * (1 + pct)
    Only touches the specific (month, bucket) cells in yoy_pct_overrides.
    """
    if projection_df is None or projection_df.empty:
        return projection_df

    month_col = projection_df.columns[0]
    df = projection_df.copy()

    for key, pct in (yoy_pct_overrides or {}).items():
        try:
            month_label, col_name = key.split("||", 1)
        except ValueError:
            continue

        if col_name not in df.columns:
            continue

        mask = df[month_col].astype(str) == str(month_label)
        if not mask.any():
            continue

        try:
            pct = float(pct)
        except Exception:
            continue

        cur = pd.to_numeric(df.loc[mask, col_name], errors="coerce").fillna(0)
        df.loc[mask, col_name] = cur * (1.0 + pct)

    # round to whole numbers for display
    for c in df.columns[1:]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).round(0).astype(int)

    return df


# -----------------------------
# Customer micro-delta overrides (SESSION)
# -----------------------------

def _cust_delta_session_key(month_label: str, col_name: str, customer_name: str) -> str:
    return f"{month_label}||{col_name}||{customer_name}"


def _get_applied_customer_deltas(request) -> dict:
    """
    Stores per-customer additive deltas (absolute units):
      { "Jan-26||CLASSIC||Falcon Farms": 20000.0, ... }
    """
    data = request.session.get("bucket_metrics_applied_customer_deltas", {})
    if not isinstance(data, dict):
        return {}

    cleaned: dict[str, float] = {}
    for k, v in data.items():
        try:
            cleaned[str(k)] = float(v)
        except Exception:
            continue
    return cleaned


def _set_applied_customer_deltas(request, overrides: dict) -> None:
    request.session["bucket_metrics_applied_customer_deltas"] = {str(k): float(v) for k, v in overrides.items()}
    request.session.modified = True


def _apply_customer_deltas_to_projection(projection_df, applied_customer_deltas: dict):
    """
    Aggregate per-customer keys to month+bucket totals and add them to the projection table.
    """
    if projection_df is None or projection_df.empty:
        return projection_df

    month_col = projection_df.columns[0]
    df = projection_df.copy()

    agg: dict[tuple[str, str], float] = {}
    for key, delta in (applied_customer_deltas or {}).items():
        try:
            month_label, col_name, _customer = key.split("||", 2)
        except ValueError:
            continue
        try:
            d = float(delta)
        except Exception:
            continue
        agg[(str(month_label), str(col_name))] = agg.get((str(month_label), str(col_name)), 0.0) + d

    for (month_label, col_name), delta_total in agg.items():
        if col_name not in df.columns:
            continue

        mask = df[month_col].astype(str) == str(month_label)
        if not mask.any():
            continue

        cur = pd.to_numeric(df.loc[mask, col_name], errors="coerce").fillna(0)
        df.loc[mask, col_name] = cur + float(delta_total)

    # round to whole numbers for display
    for c in df.columns[1:]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0).round(0).astype(int)

    return df


@login_required
@require_http_methods(["GET", "POST"])
def bucket_projections_view(request):
    tmp_path = request.session.get("bucket_metrics_tmp_path")
    if not tmp_path or not os.path.exists(tmp_path):
        return redirect("bucket_metrics")

    context = {
        "automation_name": "Bucket Projections",
        "error": None,
        "results_available": True,
        "applied_yoy": [],
        "applied_customer_deltas": [],
    }

    with open(tmp_path, "rb") as fh:
        f = BytesIO(fh.read())

    results = analyze_prognosis_workbook(f)

    reverse_map = {item["key"]: item["col"] for item in results.get("growth_fields", [])}
    request.session["bucket_metrics_growth_reverse_map"] = reverse_map
    request.session["bucket_metrics_growth_fields"] = results.get("growth_fields", [])
    request.session.modified = True

    applied_overrides = _get_applied_yoy_overrides(request)
    applied_yoy_pct = _get_applied_yoy_pct_overrides(request)
    applied_customer_deltas = _get_applied_customer_deltas(request)

    applied_growth = _get_applied_growth_overrides(request)

    # If growth edits were applied earlier, rebuild projections from that persisted growth dict
    if applied_growth:
        try:
            with open(tmp_path, "rb") as _fh:
                _bytes = _fh.read()
            projection_df_growth, yoy_suggestions_growth, start_month_label_growth = rebuild_projection_with_growth(
                BytesIO(_bytes),
                applied_growth,
            )
            results["projection_df"] = projection_df_growth
            results["yoy_suggestions"] = yoy_suggestions_growth
            results["start_month_label"] = start_month_label_growth
        except Exception:
            # If rebuild fails, fall back to baseline projections
            pass


    # Baseline = what the file said BEFORE any tweaks
    baseline_projection_df = results.get("projection_df")

    if request.method == "POST":
        action = request.POST.get("action") or ""

        # Apply: set projection cell to prev_year absolute value
        if action == "apply_yoy":
            month_label = (request.POST.get("month_label") or "").strip()
            col_name = (request.POST.get("col_name") or "").strip()
            prev_year_str = (request.POST.get("prev_year") or "").strip()

            target_value: float | None = None
            if prev_year_str:
                try:
                    target_value = float(prev_year_str)
                except Exception:
                    target_value = None

            if month_label and col_name and target_value is not None:
                key = _yoy_session_key(month_label, col_name)
                applied_overrides[key] = float(target_value)
                _set_applied_yoy_overrides(request, applied_overrides)

            return redirect("bucket_projections")

        # Unapply: remove override and any extra pct, revert to computed projection
        if action == "unapply_yoy":
            month_label = (request.POST.get("month_label") or "").strip()
            col_name = (request.POST.get("col_name") or "").strip()
            if month_label and col_name:
                key = _yoy_session_key(month_label, col_name)

                changed = False
                if key in applied_overrides:
                    applied_overrides.pop(key, None)
                    changed = True
                if key in applied_yoy_pct:
                    applied_yoy_pct.pop(key, None)
                    changed = True

                if changed:
                    _set_applied_yoy_overrides(request, applied_overrides)
                    _set_applied_yoy_pct_overrides(request, applied_yoy_pct)

            return redirect("bucket_projections")

        # Clear all YoY (and extra %)
        if action == "clear_all_yoy":
            _set_applied_yoy_overrides(request, {})
            _set_applied_yoy_pct_overrides(request, {})
            return redirect("bucket_projections")

        # Set extra % on top of an already-applied YoY cell
        if action == "set_yoy_pct":
            month_label = (request.POST.get("month_label") or "").strip()
            col_name = (request.POST.get("col_name") or "").strip()
            pct_raw = (request.POST.get("yoy_extra_pct") or "").strip()

            pct_val = 0.0
            try:
                pct_val = float(pct_raw) if pct_raw else 0.0
            except Exception:
                pct_val = 0.0

            if pct_val > 1.0:
                pct_val = pct_val / 100.0

            if month_label and col_name:
                key = _yoy_session_key(month_label, col_name)
                if key in applied_overrides:
                    if abs(pct_val) < 1e-12:
                        applied_yoy_pct.pop(key, None)
                    else:
                        applied_yoy_pct[key] = pct_val
                    _set_applied_yoy_pct_overrides(request, applied_yoy_pct)

            return redirect("bucket_projections")

        # -----------------------------
        # Customer delta apply/unapply
        # -----------------------------
        if action == "apply_customer_delta":
            month_label = (request.POST.get("month_label") or "").strip()
            col_name = (request.POST.get("col_name") or "").strip()
            customer_name = (request.POST.get("customer_name") or "").strip()
            delta_str = (request.POST.get("delta") or "").strip()

            delta_val: float | None = None
            if delta_str:
                try:
                    delta_val = float(delta_str)
                except Exception:
                    delta_val = None

            if month_label and col_name and customer_name and delta_val is not None:
                key = _cust_delta_session_key(month_label, col_name, customer_name)
                applied_customer_deltas[key] = float(delta_val)
                _set_applied_customer_deltas(request, applied_customer_deltas)

            return redirect("bucket_projections")

        # Bulk apply customer deltas (multi-select from the UI)
        if action == "apply_customer_deltas_bulk":
            items = request.POST.getlist("delta_item")
            changed = False
            for raw in items:
                # Format: month||bucket||customer||delta
                try:
                    month_label, col_name, customer_name, delta_str = raw.split("||", 3)
                except ValueError:
                    continue

                month_label = (month_label or "").strip()
                col_name = (col_name or "").strip()
                customer_name = (customer_name or "").strip()

                try:
                    delta_val = float(delta_str)
                except Exception:
                    continue

                if not (month_label and col_name and customer_name):
                    continue

                key = _cust_delta_session_key(month_label, col_name, customer_name)
                applied_customer_deltas[key] = float(delta_val)
                changed = True

            if changed:
                _set_applied_customer_deltas(request, applied_customer_deltas)

            return redirect("bucket_projections")

        if action == "unapply_customer_delta":
            month_label = (request.POST.get("month_label") or "").strip()
            col_name = (request.POST.get("col_name") or "").strip()
            customer_name = (request.POST.get("customer_name") or "").strip()

            if month_label and col_name and customer_name:
                key = _cust_delta_session_key(month_label, col_name, customer_name)
                if key in applied_customer_deltas:
                    applied_customer_deltas.pop(key, None)
                    _set_applied_customer_deltas(request, applied_customer_deltas)

            return redirect("bucket_projections")

        if action == "clear_all_customer_deltas":
            _set_applied_customer_deltas(request, {})
            return redirect("bucket_projections")

        # Growth rebuild
        if action == "apply_growth":
            growth_pct_by_safe_key = {}
            for key, val in request.POST.items():
                if key.startswith("growth__"):
                    safe_key = key.replace("growth__", "")
                    try:
                        pct = float(val) if str(val).strip() else 0.0
                    except ValueError:
                        pct = 0.0
                    growth_pct_by_safe_key[safe_key] = pct

            rev = request.session.get("bucket_metrics_growth_reverse_map", {})
            growth_real = {}
            for safe_key, pct in growth_pct_by_safe_key.items():
                real = rev.get(safe_key)
                if real:
                    growth_real[real] = pct


            # Persist growth settings so exports/ZIP rebuilds include them
            _set_applied_growth_overrides(request, growth_real)

            with open(tmp_path, "rb") as fh2:
                f2 = BytesIO(fh2.read())

            projection_df, yoy_suggestions_df, start_month_label = rebuild_projection_with_growth(f2, growth_real)

            projection_df = _apply_yoy_overrides_to_projection(projection_df, applied_overrides)
            projection_df = _apply_yoy_pct_overrides_to_projection(projection_df, applied_yoy_pct)
            projection_df = _apply_customer_deltas_to_projection(projection_df, applied_customer_deltas)

            export_path = os.path.join(tempfile.gettempdir(), f"bucket_projections_{request.user.id}.xlsx")
            with pd.ExcelWriter(export_path, engine="openpyxl") as writer:
                projection_df.to_excel(writer, index=False, sheet_name="Projections")
            request.session["bucket_metrics_projection_export_path"] = export_path
            request.session.modified = True

            results["projection_df"] = projection_df
            results["yoy_suggestions"] = yoy_suggestions_df
            results["start_month_label"] = start_month_label

    # Apply YoY replacement + extra % + customer deltas on GET too
    projection_df = results["projection_df"]
    projection_df = _apply_yoy_overrides_to_projection(projection_df, applied_overrides)
    projection_df = _apply_yoy_pct_overrides_to_projection(projection_df, applied_yoy_pct)
    projection_df = _apply_customer_deltas_to_projection(projection_df, applied_customer_deltas)

    # Export projections
    export_path = os.path.join(tempfile.gettempdir(), f"bucket_projections_{request.user.id}.xlsx")
    with pd.ExcelWriter(export_path, engine="openpyxl") as writer:
        projection_df.to_excel(writer, index=False, sheet_name="Projections")
    request.session["bucket_metrics_projection_export_path"] = export_path
    request.session.modified = True

    # Also keep a paste-ready adjustments workbook in-session so the user can
    # download it without regenerating the entire prognosis.
    try:
        adjustments_out = _generate_adjustments_to_paste_workbook(
            tmp_path=tmp_path,
            user_id=request.user.id,
            applied_customer_deltas=applied_customer_deltas,
        )
        if adjustments_out and os.path.exists(adjustments_out):
            request.session["bucket_metrics_adjustments_export_path"] = adjustments_out
            request.session.modified = True
    except Exception:
        # Don't break the page if adjustments export generation fails
        pass

    # NOTE: We intentionally do NOT regenerate the full prognosis workbook here.
    # Instead, the export ZIP includes a paste-ready "Adjustment Lines" workbook.

    # Build YoY records
    yoy_df = results.get("yoy_suggestions")
    yoy_records = []

    def _to_number_or_none(v):
        if v is None:
            return None
        try:
            if pd.isna(v):
                return None
        except Exception:
            pass
        try:
            return float(v)
        except Exception:
            return None

    if yoy_df is not None and not yoy_df.empty:
        for r in yoy_df.to_dict("records"):
            month_label = str(r.get("Month", "")).strip()
            col_name = str(r.get("Bucket Type", "")).strip()

            current_val = _to_number_or_none(r.get("This Year (current prognosis)"))
            prev_val = _to_number_or_none(r.get("Last Year (same month)"))
            pct_val = _to_number_or_none(r.get("YoY %"))

            key = _yoy_session_key(month_label, col_name)
            is_applied = key in applied_overrides

            extra_pct = float(applied_yoy_pct.get(key, 0.0))
            yoy_records.append(
                {
                    "month_label": month_label,
                    "col_name": col_name,
                    "prev_year": "" if prev_val is None else int(round(prev_val)),
                    "current_year": "" if current_val is None else int(round(current_val)),
                    "pct": "" if pct_val is None else float(pct_val),
                    "applied": is_applied,
                    "extra_pct": extra_pct,
                    "extra_pct_display": extra_pct * 100.0,
                }
            )

    applied_list = []
    for k, target in applied_overrides.items():
        try:
            m, c = k.split("||", 1)
        except ValueError:
            continue

        pct = float(applied_yoy_pct.get(k, 0.0))
        if abs(pct) > 1e-12:
            applied_list.append(f"{m} • {c} (set to {int(round(float(target)))}, +{pct * 100:.1f}%)")
        else:
            applied_list.append(f"{m} • {c} (set to {int(round(float(target)))})")

    applied_customer_list = []
    for k, delta in applied_customer_deltas.items():
        try:
            m, c, cust = k.split("||", 2)
        except ValueError:
            continue
        sign = "+" if float(delta) >= 0 else ""
        applied_customer_list.append(f"{m} • {c} • {cust} ({sign}{int(round(float(delta)))})")

    # (Optional) Customer delta suggestion records:
    # Safely renders empty if not present.
    customer_delta_records = []
    cust_df = results.get("customer_delta_suggestions")  # df from bucket_metrics.py

    def _month_sort_key(label: str):
        # "Jan-26" -> Period for correct chronological sorting
        try:
            return pd.Period(pd.to_datetime(str(label), format="%b-%y"), freq="M")
        except Exception:
            return pd.Period("1900-01", freq="M")

    if cust_df is not None and isinstance(cust_df, pd.DataFrame) and not cust_df.empty:
        df = cust_df.copy()

        # Sort by Customer first, then Month, then Bucket Type, then largest delta
        df["_month_sort"] = df["Month"].astype(str).map(_month_sort_key)
        df["_abs_delta"] = pd.to_numeric(df.get("Delta", 0), errors="coerce").fillna(0).abs()

        df = df.sort_values(
            by=["Customer", "_month_sort", "Bucket Type", "_abs_delta"],
            ascending=[True, True, True, False],
            kind="mergesort",
        )

        for r in df.to_dict("records"):
            month_label = str(r.get("Month", "")).strip()
            customer_name = str(r.get("Customer", "")).strip()
            col_name = str(r.get("Bucket Type", "")).strip()

            prev_val = _to_number_or_none(r.get("Prev Year"))
            proj_val = _to_number_or_none(r.get("Projection"))
            delta_val = _to_number_or_none(r.get("Delta"))

            key = _cust_delta_session_key(month_label, col_name, customer_name)
            is_applied = key in applied_customer_deltas

            customer_delta_records.append(
                {
                    "month_label": month_label,
                    "customer_name": customer_name,
                    "col_name": col_name,
                    "prev_year": "" if prev_val is None else int(round(prev_val)),
                    "projection": "" if proj_val is None else int(round(proj_val)),
                    "delta": "" if delta_val is None else int(round(delta_val)),
                    "applied": is_applied,
                }
            )

    context.update(
        {
            "start_month_label": results.get("start_month_label"),
            "projection_table": projection_df.to_html(classes="table table-striped table-sm", index=False, border=0),
            "growth_fields": results.get("growth_fields", []),
            "yoy_records": yoy_records,
            "applied_yoy": applied_list,
            # NEW:
            "customer_delta_records": customer_delta_records,
            "applied_customer_deltas": applied_customer_list,
        }
    )

    return render(request, "core/bucket_projections.html", context)


@login_required
def bucket_projections_export_view(request):
    tmp_path = request.session.get("bucket_metrics_tmp_path")
    if not tmp_path or not os.path.exists(tmp_path):
        return HttpResponseForbidden("Your uploaded file has expired. Please upload again.")

    export_path = request.session.get("bucket_metrics_projection_export_path")
    if not export_path or not os.path.exists(export_path):
        return HttpResponseForbidden("No export available yet. Open projections first.")

    return FileResponse(
        open(export_path, "rb"),
        as_attachment=True,
        filename="Bucket_Projections.xlsx",
    )


# Download the adjustments-only workbook (paste into Master List)
@login_required
def bucket_adjustments_export_view(request):
    export_path = request.session.get("bucket_metrics_adjustments_export_path")
    if not export_path or not os.path.exists(export_path):
        return HttpResponseForbidden("No adjustments export available yet. Open projections and apply adjustments first.")

    return FileResponse(
        open(export_path, "rb"),
        as_attachment=True,
        filename="Bucket_Adjustments_To_Paste.xlsx",
    )


@login_required
@require_http_methods(["GET", "POST"])
def run_automation(request, pk):
    automation = get_object_or_404(Automation.objects.select_related("company"), pk=pk)

    if not (request.user.is_superuser or automation.company.owner == request.user):
        return HttpResponseForbidden("You are not allowed to run this automation.")

    name_normalized = (automation.name or "").strip().lower()

    # --- Branch: Tip Tracker ---
    if name_normalized in {"tip tracker", "tips tracker", "tip tracking", "tips"}:
        return redirect("tip_tracker")

    # --- Branch: RPC Order Generator ---
    # In production the Automation.name may vary ("Retriever RPC Order", "RPC Order Generator", etc.).
    # Use a robust match so the RPC form + Order Tracker auto-add always runs.
    rpc_name_hits = {
        "retriever rpc order",
        "rpc order",
        "rpc order generator",
        "rpc generator",
        "rpc order form",
    }

    is_rpc_order_automation = (
        name_normalized in rpc_name_hits
        or (
            "rpc" in name_normalized
            and ("order" in name_normalized or "generator" in name_normalized)
        )
    )

    if is_rpc_order_automation:
        if request.method == "POST":
            form = RpcOrderForm(request.POST)
            if form.is_valid():
                files, outlook_status = generate_rpc_from_form(form.cleaned_data)

                # Auto-create/update an Order Tracker container from this RPC order
                ot_company = automation.company
                if not request.user.is_superuser:
                    owned_company = Company.objects.filter(owner=request.user).order_by("id").first()
                    if owned_company and (ot_company is None or ot_company.owner_id != request.user.id):
                        ot_company = owned_company

                try:
                    upsert_container_from_rpc_order(
                        company=ot_company,
                        created_by=request.user,
                        rpc_data=form.cleaned_data,
                    )
                except Exception as e:
                    # Don't fail the RPC download if tracking import fails
                    messages.warning(request, f"RPC generated, but Order Tracker auto-add failed: {e}")

                automation.last_run_at = timezone.now()
                automation.save(update_fields=["last_run_at"])

                first_file = files[0]
                status_text = outlook_status or "No Outlook status returned."
                messages.success(request, f"RPC generated. {status_text}")

                return FileResponse(
                    open(first_file, "rb"),
                    as_attachment=True,
                    filename=first_file.name,
                )
        else:
            form = RpcOrderForm()

        return render(request, "core/rpc_order_form.html", {"automation": automation, "form": form})

    # --- Branch: Bucket Metrics ---
    if "bucket metrics" in name_normalized:
        return redirect("bucket_metrics")

    # --- Branch: Pricing upload workflow ---
    if "pricing quote" in name_normalized or "pricing" in name_normalized:
        return redirect("pricing_upload")

    # --- Branch: RPC -> Master Sheet Formatter ---
    # The dashboard launches automations via /automations/<pk>/run/. Our Automation model
    # doesn't have a slug/route field in admin, so we dispatch by a robust name match.
    # This prevents falling through to the default BOL screen.
    is_rpc_master_formatter = (
        ("rpc" in name_normalized)
        and (
            "master" in name_normalized
            or "formatter" in name_normalized
            or "reformatter" in name_normalized
            or "master sheet" in name_normalized
        )
    )

    if is_rpc_master_formatter:
        return redirect("rpc_master_formatter")


    # --- Branch: Project Planner ---
    if "project planner" in name_normalized or "project planning" in name_normalized:
        return redirect("project_planner")

    # --- Branch: Schedule Dashboard ---
    if "schedule" in name_normalized or "calendar" in name_normalized:
        return redirect("schedule_dashboard")

    # --- Branch: Order Tracker ---
    if "order tracker" in name_normalized or "container tracker" in name_normalized or "order tracking" in name_normalized:
        return redirect("order_tracker")

    # --- Branch: Industry Relationship Web ---
    if (
        "relationship web" in name_normalized
        or "industry relationship" in name_normalized
        or "company map" in name_normalized
        or "relationship map" in name_normalized
    ):
        return redirect("industry_relationship_web")

 # --- Branch: Permaculture Garden Planner ---
    if "permaculture" in name_normalized or "garden" in name_normalized or "backyard" in name_normalized:
        return redirect("permaculture_map")

    # --- Default: BOL generator ---
    if request.method == "POST":
        form = BOLForm(request.POST)
        if form.is_valid():
            output_path = generate_bol_from_form(form.cleaned_data)

            automation.last_run_at = timezone.now()
            automation.save(update_fields=["last_run_at"])

            messages.success(request, f"Generated BOL for {automation.company.name}")

            return FileResponse(
                open(output_path, "rb"),
                as_attachment=True,
                filename=output_path.name,
            )
    else:
        form = BOLForm()

    return render(request, "core/run_bol.html", {"automation": automation, "form": form})



# -----------------------------
# Industry Relationship Web
# -----------------------------

_INDUSTRY_RELATIONSHIP_SEED = {
    "nodes": [
        {"name": "Oboya LLC", "kind": "backer", "notes": "Based in China."},
        {"name": "A-Roo", "kind": "company", "notes": "Backed by Oboya LLC."},
        {"name": "Oboya USA", "kind": "company", "notes": "Backed by Oboya LLC."},
        {"name": "Bamaplast", "kind": "supplier", "notes": "Supplier in the mapped industry web."},
        {"name": "Passion Growers", "kind": "customer", "notes": "Buys from Bamaplast."},
        {"name": "Imported Papers", "kind": "customer", "notes": "Formerly PCA."},
        {"name": "Van Tuijl", "kind": "supplier", "notes": "Supplier to Imported Papers."},
        {"name": "Plastipak", "kind": "supplier", "notes": "Supplier to Plastirex."},
        {"name": "Plastirex", "kind": "customer", "notes": "Buys from Plastipak."},
        {"name": "Naber Plastics BV", "kind": "supplier", "notes": "Supplier to Retriever Packaging and Decowraps."},
        {"name": "Retriever Packaging", "kind": "customer", "notes": "Buys from Naber Plastics BV."},
        {"name": "Decowraps", "kind": "customer", "notes": "Buys from Naber Plastics BV."},
        {"name": "EB Products & Polymers", "kind": "supplier", "notes": "Owned by Naber Plastics."},
        {"name": "Koen Pack USA", "kind": "customer", "notes": "Supplied by EB Products & Polymers."},
        {"name": "Broekhof USA", "kind": "customer", "notes": "Supplied by EB Products & Polymers."},
        {"name": "Modiform BV", "kind": "supplier", "notes": "Supplies A-Roo."},
        {"name": "Lolaar", "kind": "supplier", "notes": "Former supplier to Imported Papers."},
    ],
    "edges": [
        ("Oboya LLC", "A-Roo", "backs", False),
        ("Oboya LLC", "Oboya USA", "backs", False),
        ("Bamaplast", "Passion Growers", "supplies", False),
        ("Bamaplast", "Imported Papers", "supplies", False),
        ("Van Tuijl", "Imported Papers", "supplies", False),
        ("Plastipak", "Plastirex", "supplies", False),
        ("Naber Plastics BV", "Retriever Packaging", "supplies", False),
        ("Naber Plastics BV", "Decowraps", "supplies", False),
        ("Naber Plastics BV", "EB Products & Polymers", "owns", False),
        ("EB Products & Polymers", "Koen Pack USA", "supplies", False),
        ("EB Products & Polymers", "Broekhof USA", "supplies", False),
        ("Modiform BV", "A-Roo", "supplies", False),
        ("Lolaar", "Imported Papers", "former supplier", True),
    ],
}


def _get_user_company_or_403(request):
    if request.user.is_superuser:
        company = Company.objects.filter(name__icontains="Retriever").order_by("id").first() or Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=request.user).order_by("id").first()
    if not company:
        return None
    if not (request.user.is_superuser or company.owner_id == request.user.id):
        return None
    return company


def _seed_industry_relationship_web(company):
    if IndustryRelationshipNode.objects.filter(company=company).exists():
        return
    created = {}
    with transaction.atomic():
        for index, item in enumerate(_INDUSTRY_RELATIONSHIP_SEED["nodes"]):
            node, _ = IndustryRelationshipNode.objects.get_or_create(
                company=company,
                name=item["name"],
                defaults={
                    "kind": item.get("kind", "company"),
                    "notes": item.get("notes", ""),
                    "x": 160 + (index % 5) * 190,
                    "y": 110 + (index // 5) * 150,
                },
            )
            created[node.name] = node
        for source_name, target_name, label, is_former in _INDUSTRY_RELATIONSHIP_SEED["edges"]:
            source = created.get(source_name) or IndustryRelationshipNode.objects.get(company=company, name=source_name)
            target = created.get(target_name) or IndustryRelationshipNode.objects.get(company=company, name=target_name)
            IndustryRelationshipEdge.objects.get_or_create(
                company=company,
                source=source,
                target=target,
                label=label,
                defaults={"is_former": is_former},
            )


@login_required
def industry_relationship_web_view(request):
    company = _get_user_company_or_403(request)
    if not company:
        return HttpResponseForbidden("No company is linked to your account yet.")

    _seed_industry_relationship_web(company)

    if request.method == "POST":
        action = request.POST.get("action", "")
        if action == "add_node":
            form = IndustryRelationshipNodeForm(request.POST)
            if form.is_valid():
                node = form.save(commit=False)
                node.company = company
                node.save()
                messages.success(request, f"Added node: {node.name}")
            else:
                messages.error(request, "Could not add the node. Check for duplicates or missing fields.")
        elif action == "add_edge":
            form = IndustryRelationshipEdgeForm(request.POST, company=company)
            if form.is_valid():
                edge = form.save(commit=False)
                edge.company = company
                edge.save()
                messages.success(request, "Added relationship.")
            else:
                messages.error(request, "Could not add the relationship. Make sure both nodes are selected and are different.")
        elif action == "delete_node":
            node = get_object_or_404(IndustryRelationshipNode, pk=request.POST.get("node_id"), company=company)
            node_name = node.name
            node.delete()
            messages.success(request, f"Deleted node and its relationships: {node_name}")
        elif action == "delete_edge":
            edge = get_object_or_404(IndustryRelationshipEdge, pk=request.POST.get("edge_id"), company=company)
            edge.delete()
            messages.success(request, "Deleted relationship.")
        elif action == "reset_seed":
            IndustryRelationshipEdge.objects.filter(company=company).delete()
            IndustryRelationshipNode.objects.filter(company=company).delete()
            _seed_industry_relationship_web(company)
            messages.success(request, "Relationship web reset to the seeded map from your PDF.")
        return redirect("industry_relationship_web")

    nodes = list(IndustryRelationshipNode.objects.filter(company=company).order_by("name"))
    edges = list(
        IndustryRelationshipEdge.objects.filter(company=company)
        .select_related("source", "target")
        .order_by("source__name", "target__name")
    )

    node_form = IndustryRelationshipNodeForm()
    edge_form = IndustryRelationshipEdgeForm(company=company)

    graph_nodes = [
        {
            "id": node.id,
            "name": node.name,
            "kind": node.kind,
            "notes": node.notes,
            "x": node.x,
            "y": node.y,
        }
        for node in nodes
    ]
    graph_edges = [
        {
            "id": edge.id,
            "source": edge.source_id,
            "target": edge.target_id,
            "label": edge.label,
            "is_former": edge.is_former,
            "notes": edge.notes,
        }
        for edge in edges
    ]

    return render(
        request,
        "core/industry_relationship_web.html",
        {
            "company": company,
            "nodes": nodes,
            "edges": edges,
            "node_form": node_form,
            "edge_form": edge_form,
            "graph_json": json.dumps({"nodes": graph_nodes, "edges": graph_edges}),
            "portal_theme": get_portal_theme(user=request.user, company=company),
        },
    )


@login_required
@require_POST
def industry_relationship_positions_save_view(request):
    company = _get_user_company_or_403(request)
    if not company:
        return JsonResponse({"ok": False, "error": "Not authorized"}, status=403)
    try:
        payload = json.loads(request.body.decode("utf-8"))
        positions = payload.get("positions", [])
    except Exception:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    nodes = {n.id: n for n in IndustryRelationshipNode.objects.filter(company=company)}
    updated = 0
    for item in positions:
        try:
            node_id = int(item.get("id"))
            x = float(item.get("x"))
            y = float(item.get("y"))
        except (TypeError, ValueError):
            continue
        node = nodes.get(node_id)
        if not node:
            continue
        node.x = x
        node.y = y
        node.save(update_fields=["x", "y", "updated_at"])
        updated += 1
    return JsonResponse({"ok": True, "updated": updated})


# --- Permaculture: Excel diagram seeding ---
_PERMACULTURE_EXCEL_REL_PATH = Path(__file__).resolve().parent / "data" / "permaculture_notes.xlsx"


def _excel_argb_to_css(rgb: str) -> str:
    """Convert openpyxl ARGB/RGB to CSS hex (#RRGGBB)."""
    if not rgb:
        return ""
    rgb = str(rgb)
    if len(rgb) == 8:  # ARGB
        rgb = rgb[2:]
    if len(rgb) != 6:
        return ""
    return f"#{rgb.upper()}"


def _build_default_map_from_excel() -> tuple[int, int, dict]:
    """Parse Sheet1 of the bundled Excel notes into (rows, cols, data).

    We infer a compact used range by looking for cells that have a value, fill, or border.
    The resulting schema matches GardenMap.data version 1.
    """
    if not _PERMACULTURE_EXCEL_REL_PATH.exists():
        return 12, 18, {"version": 1, "cells": {}}

    wb = openpyxl.load_workbook(_PERMACULTURE_EXCEL_REL_PATH, data_only=True)
    if "Sheet1" in wb.sheetnames:
        ws = wb["Sheet1"]
    else:
        ws = wb.active

    max_scan_rows = min(ws.max_row or 1, 200)
    max_scan_cols = min(ws.max_column or 1, 120)

    used_max_r = 1
    used_max_c = 1

    # First pass: find the effective used range (value OR fill OR border)
    for r in range(1, max_scan_rows + 1):
        for c in range(1, max_scan_cols + 1):
            cell = ws.cell(row=r, column=c)
            val = cell.value

            has_val = val not in (None, "")
            has_fill = bool(getattr(cell.fill, "patternType", None)) and getattr(cell.fill.fgColor, "rgb", None)
            b = cell.border
            has_border = any(
                getattr(side, "style", None)
                for side in (b.left, b.right, b.top, b.bottom)
                if side is not None
            )

            if has_val or has_fill or has_border:
                used_max_r = max(used_max_r, r)
                used_max_c = max(used_max_c, c)

    rows, cols = used_max_r, used_max_c

    cells: dict[str, dict] = {}

    def infer_kind(label: str) -> str:
        s = (label or "").strip().lower()
        if not s:
            return "custom"
        if "tree" in s or s == "trees":
            return "tree"
        if "house" in s:
            return "structure"
        if "driveway" in s:
            return "path"
        if "pergola" in s:
            return "structure"
        if "water" in s and "collection" in s:
            return "water"
        # default: treat as plant if it looks like a crop name
        return "plant" if len(s) <= 40 else "custom"

    # Second pass: build cells data
    for r in range(1, rows + 1):
        for c in range(1, cols + 1):
            cell = ws.cell(row=r, column=c)
            val = cell.value

            fill_rgb = getattr(getattr(cell.fill, "fgColor", None), "rgb", None)
            css_color = _excel_argb_to_css(fill_rgb)

            b = cell.border
            border = {
                "l": getattr(b.left, "style", None) if b and b.left else None,
                "r": getattr(b.right, "style", None) if b and b.right else None,
                "t": getattr(b.top, "style", None) if b and b.top else None,
                "b": getattr(b.bottom, "style", None) if b and b.bottom else None,
            }
            has_border = any(border.values())

            label = "" if val is None else str(val).strip()

            has_val = label != ""
            has_fill = bool(css_color)

            if not (has_val or has_fill or has_border):
                continue

            kind = infer_kind(label)
            payload = {
                "kind": kind,
                "label": label,
                "notes": "",
            }
            if css_color:
                payload["color"] = css_color
            if has_border:
                payload["border"] = border

            cells[f"{r}-{c}"] = payload

    return rows, cols, {"version": 1, "cells": cells}


def _maybe_seed_permaculture_map(garden_map: GardenMap) -> bool:
    """Seed a GardenMap from the Excel diagram if it looks empty/default.

    Returns True if we modified the map.
    """
    data = garden_map.data or {}
    cells = data.get("cells") if isinstance(data, dict) else None

    # We consider a map "empty" if it has no cells saved. In that case, seed from
    # the Excel backyard diagram unless the map has already been seeded before.
    looks_empty = not cells
    seeded = isinstance(data, dict) and data.get("_seed") == "excel"

    if looks_empty and not seeded:
        rows, cols, seeded_data = _build_default_map_from_excel()
        # Mark the payload so we don't re-seed on future loads.
        if isinstance(seeded_data, dict):
            seeded_data["_seed"] = "excel"
        garden_map.rows = rows
        garden_map.cols = cols
        garden_map.data = seeded_data
        garden_map.save(update_fields=["rows", "cols", "data", "updated_at"])
        return True

    return False


@login_required
@require_http_methods(["GET", "POST"])
def permaculture_map_reset_view(request):
    """Reset the user's permaculture map back to the bundled Excel layout."""
    garden_map, _ = GardenMap.objects.get_or_create(user=request.user)
    rows, cols, data = _build_default_map_from_excel()
    if isinstance(data, dict):
        data["_seed"] = "excel"
    garden_map.rows = rows
    garden_map.cols = cols
    garden_map.data = data
    garden_map.save(update_fields=["rows", "cols", "data", "updated_at"])
    messages.success(request, "Permaculture map reset to your backyard layout.")
    return redirect("permaculture_map")


@login_required
@require_http_methods(["GET", "POST"])
def permaculture_map_import_excel_view(request):
    """Re-import the bundled Excel layout.

    Kept separate from reset so the UI can offer either action. For now this
    uses the same bundled workbook. If you later want to upload a new workbook,
    we can extend this endpoint to accept request.FILES.
    """
    garden_map, _ = GardenMap.objects.get_or_create(user=request.user)
    rows, cols, data = _build_default_map_from_excel()
    if isinstance(data, dict):
        data["_seed"] = "excel"
    garden_map.rows = rows
    garden_map.cols = cols
    garden_map.data = data
    garden_map.save(update_fields=["rows", "cols", "data", "updated_at"])
    messages.success(request, "Excel backyard layout imported.")
    return redirect("permaculture_map")


@login_required
def permaculture_map_view(request):
    # A small featured list shown immediately (the larger catalog is searched online).
    # This is also used for companion suggestions out-of-the-box.
    featured_plants = _permaculture_featured_plants()

    garden_map, _ = GardenMap.objects.get_or_create(user=request.user)

    # If this is a fresh/empty map, seed it from the bundled Excel backyard diagram.
    _maybe_seed_permaculture_map(garden_map)

    # Ensure minimum schema
    data = garden_map.data or {}
    if "version" not in data:
        data = {"version": 1, "cells": {}}
        garden_map.data = data
        garden_map.save(update_fields=["data", "updated_at"])

    context = {
        "automation_name": "Permaculture Garden Planner",
        "diagram_mode": True,
        "cell_size": 18,
        "featured_plants_json": json.dumps(featured_plants),
        "rows": garden_map.rows,
        "cols": garden_map.cols,
        "map_data_json": json.dumps(garden_map.data),
    }
    return render(request, "core/permaculture_map.html", context)


def _permaculture_featured_plants():
    """Curated starter set with richer metadata + companion relationships.

    This is NOT intended to represent "every plant"; it is a fast, built-in dataset.
    The broader catalog is queried on-demand from external sources.
    """
    return [
        {
            "name": "Basil",
            "scientific_name": "Ocimum basilicum",
            "zones": "4-10",
            "companions_good": ["Tomato", "Pepper", "Lettuce"],
            "companions_bad": ["Rue"],
            "sunlight": "Full sun",
            "water": "Moderate",
            "nitrogen": "Moderate",
            "benefits": "Repels some pests; can improve tomato flavor.",
            "drawbacks": "Sensitive to cold; bolts in heat.",
        },
        {
            "name": "Tomato",
            "scientific_name": "Solanum lycopersicum",
            "zones": "3-11",
            "companions_good": ["Basil", "Marigold", "Carrot", "Onion", "Garlic"],
            "companions_bad": ["Potato", "Fennel", "Brassicas"],
            "sunlight": "Full sun",
            "water": "Consistent",
            "nitrogen": "Moderate",
            "benefits": "High yield; trellising saves space.",
            "drawbacks": "Susceptible to blight; rotate.",
        },
        {
            "name": "Carrot",
            "scientific_name": "Daucus carota subsp. sativus",
            "zones": "3-10",
            "companions_good": ["Tomato", "Onion", "Rosemary", "Sage"],
            "companions_bad": ["Dill"],
            "sunlight": "Full sun",
            "water": "Even moisture",
            "nitrogen": "Low",
            "benefits": "Loosens soil; good edge crop.",
            "drawbacks": "Forks in rocky soil.",
        },
        {
            "name": "Marigold",
            "scientific_name": "Tagetes spp.",
            "zones": "2-11",
            "companions_good": ["Tomato", "Pepper", "Squash", "Cucumber"],
            "companions_bad": [],
            "sunlight": "Full sun",
            "water": "Low",
            "nitrogen": "Low",
            "benefits": "May reduce nematodes; attracts beneficial insects.",
            "drawbacks": "Can self-seed.",
        },
        {
            "name": "Lavender",
            "scientific_name": "Lavandula angustifolia",
            "zones": "5-9",
            "companions_good": ["Rosemary", "Sage", "Thyme"],
            "companions_bad": [],
            "sunlight": "Full sun",
            "water": "Low",
            "nitrogen": "Low",
            "benefits": "Attracts pollinators; aromatic pest deterrent.",
            "drawbacks": "Hates wet feet.",
        },
        {
            "name": "Comfrey",
            "scientific_name": "Symphytum officinale",
            "zones": "3-9",
            "companions_good": ["Apple", "Pear", "Berry bushes"],
            "companions_bad": [],
            "sunlight": "Part sun",
            "water": "Moderate",
            "nitrogen": "High",
            "benefits": "Dynamic accumulator; excellent chop-and-drop mulch.",
            "drawbacks": "Spreads if allowed to seed.",
        },
        {
            "name": "Strawberry",
            "scientific_name": "Fragaria × ananassa",
            "zones": "4-9",
            "companions_good": ["Borage", "Spinach", "Lettuce", "Bush beans"],
            "companions_bad": ["Brassicas"],
            "sunlight": "Full sun",
            "water": "Moderate",
            "nitrogen": "Moderate",
            "benefits": "Ground cover; good underplanting.",
            "drawbacks": "Bird pressure; needs mulch.",
        },
    ]


def _norm_name(s: str) -> str:
    return (s or "").strip().lower()


def _get_curated_profile_by_name(name: str) -> dict | None:
    n = _norm_name(name)
    for p in _permaculture_featured_plants():
        if _norm_name(p.get("name")) == n or _norm_name(p.get("scientific_name")) == n:
            return p
    return None


def _upsert_profile_from_curated(cur: dict) -> PlantProfile:
    sci = (cur.get("scientific_name") or cur.get("name") or "").strip()
    obj, _ = PlantProfile.objects.get_or_create(scientific_name=sci)
    obj.common_name = cur.get("name") or obj.common_name
    obj.hardiness_zones = cur.get("zones") or obj.hardiness_zones
    obj.sunlight = cur.get("sunlight") or obj.sunlight
    obj.water = cur.get("water") or obj.water
    obj.nitrogen = cur.get("nitrogen") or obj.nitrogen
    obj.benefits = cur.get("benefits") or obj.benefits
    obj.drawbacks = cur.get("drawbacks") or obj.drawbacks
    obj.companions_good = cur.get("companions_good") or obj.companions_good
    obj.companions_bad = cur.get("companions_bad") or obj.companions_bad
    obj.source = obj.source or "curated"
    obj.raw = obj.raw or {"curated": cur}
    obj.save()
    return obj


def _perenual_key() -> str | None:
    return (os.environ.get("PERENUAL_API_KEY") or "").strip() or None


@login_required
@require_http_methods(["POST"])
def permaculture_map_save_view(request):
    """Save the garden map JSON payload."""
    garden_map, _ = GardenMap.objects.get_or_create(user=request.user)
    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    # Allow optional resizing.
    rows = payload.get("rows")
    cols = payload.get("cols")
    if isinstance(rows, int) and 4 <= rows <= 60:
        garden_map.rows = rows
    if isinstance(cols, int) and 4 <= cols <= 60:
        garden_map.cols = cols

    data = payload.get("data")
    if not isinstance(data, dict):
        return JsonResponse({"ok": False, "error": "Missing data"}, status=400)

    # Basic shape validation.
    if data.get("version") != 1 or not isinstance(data.get("cells", {}), dict):
        return JsonResponse({"ok": False, "error": "Unsupported schema"}, status=400)

    garden_map.data = data
    garden_map.save(update_fields=["rows", "cols", "data", "updated_at"])
    return JsonResponse({"ok": True})


@login_required
@require_http_methods(["GET"])
def permaculture_plant_search_view(request):
    """Search a large online plant catalog and return normalized results.

    Priority order:
      1) Perenual (if PERENUAL_API_KEY is configured) - richer gardening metadata
      2) GBIF (no key) - massive taxonomy, but minimal gardening metadata
    """
    q = (request.GET.get("q") or "").strip()
    if len(q) < 2:
        return JsonResponse({"ok": True, "results": []})

    perenual_key = (os.environ.get("PERENUAL_API_KEY") or "").strip()
    if perenual_key:
        # Perenual species list search (requires API key)
        url = "https://perenual.com/api/v2/species-list?" + urllib.parse.urlencode(
            {"key": perenual_key, "q": q, "page": 1}
        )
        try:
            with urllib.request.urlopen(url, timeout=12) as resp:
                raw = resp.read().decode("utf-8")
            data = json.loads(raw)
        except Exception:
            return JsonResponse({"ok": False, "error": "Plant search failed"}, status=502)

        results = []
        for item in (data.get("data") or [])[:20]:
            sci_list = item.get("scientific_name") or []
            sci = (sci_list[0] if isinstance(sci_list, list) and sci_list else "")
            results.append(
                {
                    "scientific_name": sci,
                    "common_name": item.get("common_name") or "",
                    "rank": "species",
                    "family": item.get("family") or "",
                    "genus": item.get("genus") or "",
                    "species": item.get("species_epithet") or "",
                    "perenual_id": item.get("id"),
                }
            )
        return JsonResponse({"ok": True, "provider": "perenual", "results": results})

    # Fallback: GBIF taxonomy search (no key)
    params = {"q": q, "kingdom": "Plantae", "limit": 20}
    url = "https://api.gbif.org/v1/species/search?" + urllib.parse.urlencode(params)

    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            raw = resp.read().decode("utf-8")
        gbif = json.loads(raw)
    except Exception:
        return JsonResponse({"ok": False, "error": "Plant search failed"}, status=502)

    results = []
    for item in gbif.get("results", [])[:20]:
        sci = item.get("scientificName") or item.get("canonicalName")
        if not sci:
            continue
        results.append(
            {
                "scientific_name": sci,
                "common_name": (item.get("vernacularName") or ""),
                "rank": (item.get("rank") or ""),
                "family": (item.get("family") or ""),
                "genus": (item.get("genus") or ""),
                "species": (item.get("species") or ""),
                "gbif_key": item.get("key"),
            }
        )

    return JsonResponse({"ok": True, "provider": "gbif", "results": results})


@login_required
@require_http_methods(["GET"])
def permaculture_plant_profile_view(request):
    """Return a normalized plant profile.

    If PERENUAL_API_KEY is configured and a perenual_id is provided, we'll fetch
    plant details and cache them in PlantProfile.
    """
    perenual_id = request.GET.get("perenual_id")
    scientific_name = (request.GET.get("scientific_name") or "").strip()

    # 1) If we have an exact cached profile, return it.
    if scientific_name:
        cached = PlantProfile.objects.filter(scientific_name=scientific_name).first()
        if cached:
            return JsonResponse(
                {
                    "ok": True,
                    "profile": {
                        "scientific_name": cached.scientific_name,
                        "common_name": cached.common_name,
                        "hardiness_zones": cached.hardiness_zones,
                        "sunlight": cached.sunlight,
                        "water": cached.water,
                        "nitrogen": cached.nitrogen,
                        "benefits": cached.benefits,
                        "drawbacks": cached.drawbacks,
                        "companions_good": cached.companions_good,
                        "companions_bad": cached.companions_bad,
                        "source": cached.source,
                    },
                }
            )

    # 2) If the plant is in our curated starter set, upsert into cache.
    curated = _get_curated_profile_by_name(scientific_name) if scientific_name else None
    if curated:
        obj = _upsert_profile_from_curated(curated)
        return JsonResponse(
            {
                "ok": True,
                "profile": {
                    "scientific_name": obj.scientific_name,
                    "common_name": obj.common_name,
                    "hardiness_zones": obj.hardiness_zones,
                    "sunlight": obj.sunlight,
                    "water": obj.water,
                    "nitrogen": obj.nitrogen,
                    "benefits": obj.benefits,
                    "drawbacks": obj.drawbacks,
                    "companions_good": obj.companions_good,
                    "companions_bad": obj.companions_bad,
                    "source": obj.source,
                },
            }
        )

    # 3) If we have a Perenual id + key, fetch details and cache.
    key = _perenual_key()
    if key and perenual_id:
        url = f"https://perenual.com/api/v2/species/details/{urllib.parse.quote(str(perenual_id))}?" + urllib.parse.urlencode(
            {"key": key}
        )
        try:
            with urllib.request.urlopen(url, timeout=12) as resp:
                raw = resp.read().decode("utf-8")
            details = json.loads(raw)
        except Exception:
            return JsonResponse({"ok": False, "error": "Plant profile fetch failed"}, status=502)

        sci_list = details.get("scientific_name") or []
        sci = (sci_list[0] if isinstance(sci_list, list) and sci_list else "").strip()
        if not sci:
            sci = scientific_name or f"Perenual:{perenual_id}"

        obj, _ = PlantProfile.objects.get_or_create(scientific_name=sci)
        obj.common_name = details.get("common_name") or obj.common_name
        hardiness = details.get("hardiness") or {}
        if isinstance(hardiness, dict):
            mn = hardiness.get("min")
            mx = hardiness.get("max")
            if mn and mx:
                obj.hardiness_zones = f"{mn}-{mx}"
        obj.sunlight = ", ".join(details.get("sunlight") or []) if isinstance(details.get("sunlight"), list) else (details.get("sunlight") or obj.sunlight)
        obj.water = details.get("watering") or obj.water
        obj.source = "perenual"
        obj.raw = details
        obj.save()

        return JsonResponse(
            {
                "ok": True,
                "profile": {
                    "scientific_name": obj.scientific_name,
                    "common_name": obj.common_name,
                    "hardiness_zones": obj.hardiness_zones,
                    "sunlight": obj.sunlight,
                    "water": obj.water,
                    "nitrogen": obj.nitrogen,
                    "benefits": obj.benefits,
                    "drawbacks": obj.drawbacks,
                    "companions_good": obj.companions_good,
                    "companions_bad": obj.companions_bad,
                    "source": obj.source,
                },
            }
        )

    # Last resort: return minimal profile (user can still place it)
    return JsonResponse(
        {
            "ok": True,
            "profile": {
                "scientific_name": scientific_name or "",
                "common_name": "",
                "hardiness_zones": "",
                "sunlight": "",
                "water": "",
                "nitrogen": "",
                "benefits": "",
                "drawbacks": "",
                "companions_good": [],
                "companions_bad": [],
                "source": "unknown",
            },
        }
    )


@login_required
@require_http_methods(["GET"])
def permaculture_companion_suggest_view(request):
    """Suggest companion plants for a given plant, considering nearby placements."""
    plant = (request.GET.get("plant") or "").strip()
    nearby_raw = (request.GET.get("nearby") or "").strip()
    nearby = [p.strip() for p in nearby_raw.split(",") if p.strip()]

    # Ensure curated plants exist in cache so suggestions work immediately.
    for cur in _permaculture_featured_plants():
        _upsert_profile_from_curated(cur)

    # Load profile by either scientific or common name
    prof = PlantProfile.objects.filter(scientific_name__iexact=plant).first()
    if not prof:
        prof = PlantProfile.objects.filter(common_name__iexact=plant).first()

    good = list(prof.companions_good) if prof else []
    bad = list(prof.companions_bad) if prof else []

    nearby_norm = {_norm_name(x) for x in nearby}

    recommended = [x for x in good if _norm_name(x) not in nearby_norm]
    conflicts = [x for x in nearby if _norm_name(x) in {_norm_name(b) for b in bad}]

    return JsonResponse(
        {
            "ok": True,
            "plant": plant,
            "recommended": recommended[:12],
            "conflicts": conflicts[:12],
            "known": bool(prof),
        }
    )


def project_planner_view(request):
    """
    Project Planner:
      - Create + edit project ideas with cost/time/difficulty/risk/priority
      - Sort by priority/cost/time/difficulty/risk/name/created/updated
    """
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        try:
            company = Company.objects.get(owner=user)
        except Company.DoesNotExist:
            company = None

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    # Sorting
    sort_key = (request.GET.get("sort") or "priority").strip().lower()
    sort_dir = (request.GET.get("dir") or "desc").strip().lower()

    sort_map = {
        "priority": "priority_level",
        "priority_level": "priority_level",
        "cost": "estimated_cost",
        "estimated_cost": "estimated_cost",
        "time": "estimated_time_hours",
        "hours": "estimated_time_hours",
        "estimated_time_hours": "estimated_time_hours",
        "difficulty": "estimated_difficulty",
        "estimated_difficulty": "estimated_difficulty",
        "risk": "risk_factor",
        "risk_factor": "risk_factor",
        "name": "project_name",
        "project_name": "project_name",
        "created": "created_at",
        "created_at": "created_at",
        "updated": "updated_at",
        "updated_at": "updated_at",
    }

    field = sort_map.get(sort_key, "priority_level")
    prefix = "-" if sort_dir != "asc" else ""
    ordering = [f"{prefix}{field}", "-updated_at", "-created_at"]

    projects = ProjectPlanEntry.objects.filter(company=company, completed=False).order_by(*ordering)
    completed_projects = ProjectPlanEntry.objects.filter(company=company, completed=True).order_by('-completed_at', '-updated_at')

    # Load edit instance (via ?edit=<id>)
    edit_id = request.GET.get("edit")
    instance = None
    if edit_id:
        instance = get_object_or_404(ProjectPlanEntry, id=edit_id, company=company)

    if request.method == "POST":
        entry_id = request.POST.get("entry_id") or None
        if entry_id:
            instance = get_object_or_404(ProjectPlanEntry, id=entry_id, company=company)

        form = ProjectPlanEntryForm(request.POST, instance=instance)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.company = company
            obj.user = user
            obj.save()
            messages.success(request, "Project saved.")
            return redirect("project_planner")
    else:
        form = ProjectPlanEntryForm(instance=instance)
        if instance:
            form.initial["entry_id"] = instance.id

    context = {
        "automation_name": "Project Planner",
        "company": company,
        "projects": projects,
        "completed_projects": completed_projects,
        "form": form,
        "sort": sort_key,
        "dir": sort_dir,
        "edit_instance": instance,
    }
    return render(request, "core/project_planner.html", context)


@login_required
@require_http_methods(["GET", "POST"])


@login_required
@require_http_methods(["POST"])
def project_plan_complete_view(request, pk: int):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        try:
            company = Company.objects.get(owner=user)
        except Company.DoesNotExist:
            company = None

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    project = get_object_or_404(ProjectPlanEntry, id=pk, company=company)

    if not project.completed:
        project.completed = True
        project.completed_at = timezone.now()
        project.save(update_fields=["completed", "completed_at", "updated_at"])
        messages.success(request, f"Completed: {project.project_name}")
    else:
        messages.info(request, f"Already completed: {project.project_name}")

    return redirect("project_planner")


def project_plan_delete_view(request, pk: int):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        try:
            company = Company.objects.get(owner=user)
        except Company.DoesNotExist:
            company = None

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    project = get_object_or_404(ProjectPlanEntry, id=pk, company=company)

    if request.method == "POST":
        project.delete()
        messages.success(request, "Project deleted.")
        return redirect("project_planner")

    return render(request, "core/project_plan_confirm_delete.html", {"project": project, "automation_name": "Project Planner"})


def bucket_metrics_view(request, automation_id=None):
    context = {
        "automation_name": "Bucket Metrics from Prognosis Spreadsheet",
        "results_available": False,
    }

    # Step 2: Apply growth (no re-upload)
    if request.method == "POST" and request.POST.get("action") == "apply_growth":
        tmp_path = request.session.get("bucket_metrics_tmp_path")

        if not tmp_path or not os.path.exists(tmp_path):
            context["error"] = "Your uploaded file has expired. Please upload the spreadsheet again."
            return render(request, "core/bucket_metrics.html", context)

        growth_pct_by_safe_key = {}
        for key, val in request.POST.items():
            if key.startswith("growth__"):
                safe_key = key.replace("growth__", "")
                try:
                    pct = float(val) if str(val).strip() else 0.0
                except ValueError:
                    pct = 0.0
                growth_pct_by_safe_key[safe_key] = pct

        reverse_map = request.session.get("bucket_metrics_growth_reverse_map", {})
        growth_real = {}
        for safe_key, pct in growth_pct_by_safe_key.items():
            real = reverse_map.get(safe_key)
            if real:
                growth_real[real] = pct


        # Persist growth settings so exports/ZIP rebuilds include them
        _set_applied_growth_overrides(request, growth_real)

        try:
            with open(tmp_path, "rb") as fh:
                f = BytesIO(fh.read())

            projection_df, yoy_suggestions, start_month_label = rebuild_projection_with_growth(f, growth_real)

            context.update(
                {
                    "results_available": True,
                    "start_month_label": start_month_label,
                    "projection_table": projection_df.to_html(classes="table table-striped table-sm", index=False, border=0),
                    "yoy_suggestions_table": yoy_suggestions.to_html(classes="table table-striped table-sm", index=False, border=0),
                    "growth_fields": request.session.get("bucket_metrics_growth_fields", []),
                    "top_customers_table": request.session.get("bucket_metrics_top_customers_table"),
                    "per_customer_month_table": request.session.get("bucket_metrics_per_customer_month_table"),
                    "per_customer_city_item_table": request.session.get("bucket_metrics_per_customer_city_item_table"),
                    "per_customer_city_item_month_table": request.session.get("bucket_metrics_per_customer_city_item_month_table"),
                }
            )
        except Exception as e:
            context["error"] = f"Error rebuilding projections: {e}"

        context["projections_url"] = "bucket_projections"
        return render(request, "core/bucket_metrics.html", context)

    # Step 1: Initial upload
    if request.method == "POST" and request.FILES.get("file"):
        excel_file = request.FILES["file"]

        try:
            suffix = os.path.splitext(excel_file.name or "")[1] or ".xlsx"

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                for chunk in excel_file.chunks():
                    tmp.write(chunk)
                tmp_path = tmp.name

            request.session["bucket_metrics_tmp_path"] = tmp_path
            request.session["bucket_metrics_uploaded_name"] = excel_file.name

            # New upload: clear persisted growth edits
            _set_applied_growth_overrides(request, {})

            with open(tmp_path, "rb") as fh:
                f = BytesIO(fh.read())

            results = analyze_prognosis_workbook(f)

            reverse_map = {item["key"]: item["col"] for item in results.get("growth_fields", [])}
            request.session["bucket_metrics_growth_reverse_map"] = reverse_map
            request.session["bucket_metrics_growth_fields"] = results.get("growth_fields", [])

            top_customers_html = results["top_customers"].to_html(classes="table table-striped table-sm", index=False, border=0)
            per_customer_month_html = results["per_customer_month"].to_html(classes="table table-striped table-sm", index=False, border=0)
            per_customer_city_item_html = results["per_customer_city_item"].to_html(classes="table table-striped table-sm", index=False, border=0)
            per_customer_city_item_month_html = results["per_customer_city_item_month"].to_html(classes="table table-striped table-sm", index=False, border=0)

            request.session["bucket_metrics_top_customers_table"] = top_customers_html
            request.session["bucket_metrics_per_customer_month_table"] = per_customer_month_html
            request.session["bucket_metrics_per_customer_city_item_table"] = per_customer_city_item_html
            request.session["bucket_metrics_per_customer_city_item_month_table"] = per_customer_city_item_month_html
            request.session.modified = True

            context.update(
                {
                    "results_available": True,
                    "start_month_label": results.get("start_month_label"),
                    "top_customers_table": top_customers_html,
                    "per_customer_month_table": per_customer_month_html,
                    "per_customer_city_item_table": per_customer_city_item_html,
                    "per_customer_city_item_month_table": per_customer_city_item_month_html,
                    "projection_table": results["projection_df"].to_html(classes="table table-striped table-sm", index=False, border=0),
                    "yoy_suggestions_table": results["yoy_suggestions"].to_html(classes="table table-striped table-sm", index=False, border=0),
                    "growth_fields": results.get("growth_fields", []),
                }
            )
        except Exception as e:
            context["error"] = f"Error reading file: {e}"

        return render(request, "core/bucket_metrics.html", context)

    return render(request, "core/bucket_metrics.html", context)


def _get_company_for_request(request):
    user = request.user

    if user.is_superuser:
        company_id = request.GET.get("company_id")
        if company_id:
            return get_object_or_404(Company, id=company_id)
        return Company.objects.order_by("id").first()

    try:
        return Company.objects.get(owner=user)
    except Company.DoesNotExist:
        return None


# -----------------------------
# Pricing workflow
# -----------------------------

@login_required
@require_http_methods(["GET", "POST"])
def pricing_upload_view(request):
    company = _get_company_for_request(request)
    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    if request.method == "POST":
        form = PricingUploadForm(request.POST, request.FILES)
        if form.is_valid():
            f = form.cleaned_data["file"]

            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
                for chunk in f.chunks():
                    tmp.write(chunk)
                tmp_path = tmp.name

            rows = parse_pricing_matrix_csv(tmp_path)

            existing_meta = {}
            for line in PricingQuoteLine.objects.select_related("customer").filter(company=company):
                key = (
                    (line.customer.name or "").strip(),
                    (line.destination or "").strip(),
                    (line.product_description or "").strip(),
                )
                existing_meta[key] = {
                    "pallet_quantity_pieces": line.pallet_quantity_pieces,
                    "include_in_quote": line.include_in_quote,
                }

            PricingQuoteLine.objects.filter(company=company).delete()

            normalized_rows = {}
            for r in rows:
                canon_customer = normalize_customer_name(r.get("customer", ""))
                if canon_customer is None:
                    continue

                dest_raw = r.get("destination", "")
                prod = (r.get("product_description") or "").strip()
                price = r["price_delivered"]

                norm_dest = normalize_destination(canon_customer, dest_raw)

                key = (canon_customer, norm_dest, prod)
                normalized_rows[key] = price

            created_lines = 0
            updated_lines = 0
            created_customers = 0

            for (canon_customer, norm_dest, prod), price in normalized_rows.items():
                customer_obj, cust_created = get_or_create_customer_safe(company, canon_customer)
                if cust_created:
                    created_customers += 1

                meta = existing_meta.get((canon_customer, norm_dest, prod), None)
                saved_pallet_qty = (meta or {}).get("pallet_quantity_pieces", 0)
                default_pallet_qty = get_rpc_default_pallet_quantity_pieces(prod)
                pallet_qty = int(saved_pallet_qty or default_pallet_qty or 0)
                include = meta["include_in_quote"] if meta else True

                PricingQuoteLine.objects.create(
                    company=company,
                    customer=customer_obj,
                    destination=norm_dest,
                    product_description=prod,
                    price_delivered=price,
                    pallet_quantity_pieces=pallet_qty,
                    include_in_quote=include,
                )

                created_lines += 1

            merge_duplicate_pricing_customers(company)

            messages.success(
                request,
                f"Imported pricing: {created_customers} customers, {created_lines} new lines, {updated_lines} updated lines.",
            )
            return redirect("pricing_customer_list")

    else:
        form = PricingUploadForm()

    return render(request, "core/pricing_upload.html", {"form": form, "company": company})


@login_required
def pricing_customer_list_view(request):
    company = _get_company_for_request(request)
    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    customers = PricingCustomer.objects.filter(company=company).order_by("name")

    return render(
        request,
        "core/pricing_customer_list.html",
        {
            "company": company,
            "customers": customers,
        },
    )


@login_required
@require_http_methods(["GET", "POST"])
def pricing_customer_edit_view(request, customer_id):
    """
    Permanent edits (saved): pallet_quantity_pieces, include_in_quote
    Temporary quote-only edits (session only): product description override
    """
    company = _get_company_for_request(request)
    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    customer = get_object_or_404(PricingCustomer, id=customer_id, company=company)

    lines_qs = PricingQuoteLine.objects.filter(company=company, customer=customer).order_by(
        "destination", "product_description"
    )

    currency_code, currency_symbol = get_currency_for_customer_name(customer.name)

    overrides = get_quote_desc_overrides(request, company.id, customer.id)

    if request.method == "POST":
        if request.POST.get("clear_quote_desc") == "1":
            clear_quote_desc_overrides(request, company.id, customer.id)
            messages.success(request, "Cleared quote-only description edits.")
            return redirect("pricing_customer_edit", customer_id=customer.id)

        new_overrides = dict(overrides)

        for line in lines_qs:
            qty_key = f"pallet_{line.id}"
            if qty_key in request.POST:
                raw = (request.POST.get(qty_key) or "").strip()
                try:
                    line.pallet_quantity_pieces = int(raw) if raw else 0
                except ValueError:
                    pass

            include_key = f"include_{line.id}"
            line.include_in_quote = include_key in request.POST

            line.save(update_fields=["pallet_quantity_pieces", "include_in_quote"])

            desc_key = f"quote_desc_{line.id}"
            raw_desc = (request.POST.get(desc_key) or "").strip()

            if raw_desc and raw_desc != line.product_description:
                new_overrides[str(line.id)] = raw_desc
            else:
                new_overrides.pop(str(line.id), None)

        bulk_exclude_product = (request.POST.get("bulk_exclude_product") or "").strip()
        if bulk_exclude_product:
            PricingQuoteLine.objects.filter(
                company=company,
                customer=customer,
                product_description=bulk_exclude_product,
            ).update(include_in_quote=False)
            messages.success(
                request,
                f'Deselected "{bulk_exclude_product}" for all locations for this customer.',
            )
        else:
            messages.success(
                request,
                "Saved pallet quantities / inclusions. Quote-only descriptions updated for the next quote.",
            )

        set_quote_desc_overrides(request, company.id, customer.id, new_overrides)
        return redirect("pricing_customer_edit", customer_id=customer.id)

    for line in lines_qs:
        line.quote_desc_value = overrides.get(str(line.id), line.product_description)

    return render(
        request,
        "core/pricing_customer_edit.html",
        {
            "company": company,
            "customer": customer,
            "lines": lines_qs,
            "currency_code": currency_code,
            "currency_symbol": currency_symbol,
            "has_quote_desc_overrides": bool(overrides),
        },
    )


@login_required
def pricing_customer_quote_view(request, customer_id):
    """
    HTML quote page (easy to print/save as PDF).
    Uses quote-only description overrides ONCE, then clears them.
    """
    company = _get_company_for_request(request)
    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    customer = get_object_or_404(PricingCustomer, id=customer_id, company=company)
    # NOTE: We materialize to a list because we compute a pivot grid in-memory for the quote.
    lines = list(
        PricingQuoteLine.objects.filter(
            company=company,
            customer=customer,
            include_in_quote=True,
        ).order_by("destination", "product_description")
    )

    currency_code, currency_symbol = get_currency_for_customer_name(customer.name)

    overrides = get_quote_desc_overrides(request, company.id, customer.id)

    for line in lines:
        line.display_product_description = overrides.get(
            str(line.id),
            line.product_description,
        )

    # Build pivot axes
    # X axis: destinations (locations)
    destinations: list[str] = []
    seen_dests: set[str] = set()

    # Y axis: products are keyed by (description, pallet_quantity_pieces)
    # so the generated sheet keeps separate rows when the same description
    # appears with different pallet quantities.
    row_keys: list[tuple[str, int]] = []
    seen_row_keys: set[tuple[str, int]] = set()
    qtys_by_product: dict[str, set[int]] = {}

    for line in lines:
        d = (line.destination or "").strip() or "(Unspecified)"
        p = (line.display_product_description or "").strip() or "(Unspecified)"
        qty = int(line.pallet_quantity_pieces or 0)
        row_key = (p, qty)

        if d not in seen_dests:
            destinations.append(d)
            seen_dests.add(d)

        if row_key not in seen_row_keys:
            row_keys.append(row_key)
            seen_row_keys.add(row_key)

        qtys_by_product.setdefault(p, set()).add(qty)

    # grid[(product, qty)][destination] = PricingQuoteLine | None
    grid: dict[tuple[str, int], dict[str, PricingQuoteLine | None]] = {
        row_key: {d: None for d in destinations} for row_key in row_keys
    }
    for line in lines:
        d = (line.destination or "").strip() or "(Unspecified)"
        p = (line.display_product_description or "").strip() or "(Unspecified)"
        qty = int(line.pallet_quantity_pieces or 0)
        row_key = (p, qty)
        # If duplicates exist, keep the first one (stable with queryset ordering)
        if row_key in grid and d in grid[row_key] and grid[row_key][d] is None:
            grid[row_key][d] = line

    def _display_product_label(product_name: str, qty: int) -> str:
        if len(qtys_by_product.get(product_name, set())) > 1:
            return f"{product_name} ({qty} pcs / pallet)"
        return product_name

    # Precompute a template-friendly structure: one row per product/qty combination
    def _is_grey_product(product_name: str) -> bool:
        return "grey" in (product_name or "").strip().lower()

    row_keys_non_grey = [rk for rk in row_keys if not _is_grey_product(rk[0])]
    row_keys_grey = [rk for rk in row_keys if _is_grey_product(rk[0])]

    quote_rows = [
        {
            "product": _display_product_label(product_name, qty),
            "cells": [grid[(product_name, qty)].get(d) for d in destinations],
        }
        for (product_name, qty) in row_keys_non_grey
    ]

    grey_quote_rows = [
        {
            "product": _display_product_label(product_name, qty),
            "cells": [grid[(product_name, qty)].get(d) for d in destinations],
        }
        for (product_name, qty) in row_keys_grey
    ]

    if overrides:
        clear_quote_desc_overrides(request, company.id, customer.id)

    return render(
        request,
        "core/pricing_quote.html",
        {
            "company": company,
            "customer": customer,
            "lines": lines,
            "destinations": destinations,
            "quote_rows": quote_rows,
            "grey_quote_rows": grey_quote_rows,
            "currency_code": currency_code,
            "currency_symbol": currency_symbol,
        },
    )

@require_POST
@login_required
def order_container_delete_view(request, container_id: int):
    user = request.user

    if user.is_superuser:
        # Superusers can delete any container by ID
        container = get_object_or_404(OrderContainer, id=container_id)
    else:
        # Normal users are scoped to their Company via Company.owner
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, id=container_id, company=company)

    # This will cascade-delete lines + documents if your FKs are CASCADE (they should be).
    container.delete()

    messages.success(request, "Order deleted.")
    return redirect("order_tracker")


@require_POST
@login_required
def order_container_toggle_delivered_view(request, container_id: int):
    """Toggle an order between Active and Delivered.

    Delivered is represented by status == 'Delivered' (case-insensitive).
    When toggling from Delivered -> Active, we set status to 'In transit' (safe default).
    """
    user = request.user

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, id=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, id=container_id, company=company)

    if (container.status or "").strip().lower() == "delivered":
        container.status = "In transit"
        container.save(update_fields=["status", "updated_at"])
        messages.success(request, "Order moved back to Active.")
    else:
        container.status = "Delivered"
        container.save(update_fields=["status", "updated_at"])
        messages.success(request, "Order marked Delivered.")

    # Return to where the user was (keep filters), fall back to tracker.
    return redirect(request.META.get("HTTP_REFERER") or "order_tracker")
    
@login_required
def order_tracker_view(request):
    """
    Sea container / order tracking dashboard.
    Shows containers scoped to the user's company (owner) and allows create/edit.
    """
    user = request.user

    # Superusers see ALL containers across companies.
    if user.is_superuser:
        company = None
        containers = OrderContainer.objects.all().prefetch_related("tags")
    else:
        company = get_object_or_404(Company, owner=user)
        containers = OrderContainer.objects.filter(company=company).prefetch_related("tags")

    # Archived containers should not appear on the main tracker.
    containers = containers.filter(is_archived=False)

    q = (request.GET.get("q") or "").strip()
    status = (request.GET.get("status") or "").strip()
    assigned_to = (request.GET.get("assigned_to") or "").strip()

    sort = (request.GET.get("sort") or "").strip()
    direction = (request.GET.get("dir") or "desc").strip().lower()

    if q:
        containers = containers.filter(
            Q(customer_name__icontains=q)
            | Q(location_name__icontains=q)
            | Q(po_number__icontains=q)
            | Q(rpc_number__icontains=q)
            | Q(container_number__icontains=q)
            | Q(carrier__icontains=q)
            | Q(booking_number__icontains=q)
            | Q(bill_of_lading_number__icontains=q)
        )

    if status:
        # status is free text; treat filter as "contains" so you can search partial statuses
        containers = containers.filter(status__icontains=status)

    if assigned_to:
        containers = containers.filter(assigned_to__iexact=assigned_to)

    # Sorting (whitelisted)
    sort_map = {
        "customer": "customer_name",
        "location": "location_name",
        "requested": "requested_date",
        "eta": "eta",
        "est_delivery": "estimated_delivery_date",
        "loading": "loading_date",
        "assigned_to": "assigned_to",
    }
    sort_field = sort_map.get(sort)

    if direction not in {"asc", "desc"}:
        direction = "desc"

    if sort_field:
        prefix = "" if direction == "asc" else "-"
        containers = containers.order_by(f"{prefix}{sort_field}", "-updated_at", "-created_at")
    else:
        containers = containers.order_by("-updated_at", "-created_at")

    # Split active vs delivered (delivered means status == "Delivered" case-insensitive)
    delivered_q = Q(status__iexact="delivered")
    delivered_containers = list(containers.filter(delivered_q))
    active_containers = list(containers.exclude(delivered_q))

    # Attach JSONCargo updates to each row so the master tracker can show the
    # latest API ETA next to the manually maintained ETA. We keep both values:
    # - latest_pending_tracking_update: actionable pending change/error
    # - latest_tracking_update: most recent JSONCargo result, even after approval
    # This prevents the API ETA column from disappearing after you acknowledge
    # an update.
    try:
        from core.models import OrderContainerTrackingUpdate

        visible_ids = [c.id for c in active_containers] + [c.id for c in delivered_containers]

        latest_pending_by_container = {}
        for upd in OrderContainerTrackingUpdate.objects.filter(
            container_id__in=visible_ids,
            status=OrderContainerTrackingUpdate.STATUS_PENDING,
        ).order_by("-created_at", "-id"):
            latest_pending_by_container.setdefault(upd.container_id, upd)

        latest_any_by_container = {}
        for upd in OrderContainerTrackingUpdate.objects.filter(
            container_id__in=visible_ids,
        ).order_by("-created_at", "-id"):
            latest_any_by_container.setdefault(upd.container_id, upd)

        pending_container_ids = set(latest_pending_by_container.keys())
        for c in active_containers + delivered_containers:
            c.latest_pending_tracking_update = latest_pending_by_container.get(c.id)
            c.latest_tracking_update = latest_any_by_container.get(c.id)
    except Exception:
        pending_container_ids = set()
        for c in active_containers + delivered_containers:
            c.latest_pending_tracking_update = None
            c.latest_tracking_update = None

    # --- Vessel map data (MyShipTracking) ---
    vessel_points = []
    vessel_map_error = None
    try:
        mmsi_list = [int(c.vessel_mmsi) for c in active_containers if getattr(c, "vessel_mmsi", None)]
        imo_list = [int(c.vessel_imo) for c in active_containers if getattr(c, "vessel_imo", None)]
    except Exception:
        mmsi_list, imo_list = [], []

    if mmsi_list or imo_list:
        data, err = bulk_vessel_status(mmsi_list=mmsi_list, imo_list=imo_list)
        vessel_map_error = err
        if data:
            # Index returned vessel data by identifiers for quick matching to containers.
            by_mmsi = {str(v.get("mmsi")): v for v in data if v.get("mmsi") is not None}
            by_imo = {str(v.get("imo")): v for v in data if v.get("imo") is not None}

            for c in active_containers:
                v = None
                if getattr(c, "vessel_mmsi", None):
                    v = by_mmsi.get(str(c.vessel_mmsi))
                if v is None and getattr(c, "vessel_imo", None):
                    v = by_imo.get(str(c.vessel_imo))
                if not v:
                    continue

                try:
                    lat = float(v.get("lat"))
                    lng = float(v.get("lng"))
                except Exception:
                    continue

                vessel_points.append(
                    {
                        "lat": lat,
                        "lng": lng,
                        "vessel_name": (getattr(c, "vessel_name", "") or v.get("vessel_name") or "").strip(),
                        "mmsi": v.get("mmsi"),
                        "imo": v.get("imo"),
                        "received": v.get("received"),
                        "container_label": f"{c.customer_name}{(' • ' + c.location_name) if c.location_name else ''}",
                        "container_url": reverse("order_container_edit", kwargs={"container_id": c.id}),
                    }
                )

    return render(
        request,
        "core/order_tracker.html",
        {
            "automation_name": "Order Tracker",
            "company": company,
            "company_display": (company.name if company else "All Companies"),
            "containers": active_containers,
            "delivered_containers": delivered_containers,
            "q": q,
            "status": status,
            "assigned_to": assigned_to,
            "sort": sort,
            "dir": direction,
            "vessel_points_json": json.dumps(vessel_points),
            "vessel_map_error": vessel_map_error,
            "pending_container_ids": pending_container_ids,
        },
    )


@require_POST
@login_required
def order_container_archive_view(request, container_id: int):
    """Archive an order so it no longer appears in the tracker or JSONCargo syncing."""
    user = request.user

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, id=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, id=container_id, company=company)

    container.is_archived = True
    container.archived_at = timezone.now()
    # Safe default: delivered orders are the usual archive target.
    if (container.status or "").strip() == "":
        container.status = "Delivered"
    container.save(update_fields=["is_archived", "archived_at", "status", "updated_at"])

    messages.success(request, "Order archived. It will no longer sync or show tracking updates.")
    return redirect(request.META.get("HTTP_REFERER") or "order_tracker")


@require_POST
@login_required
def order_container_unarchive_view(request, container_id: int):
    """Un-archive an order (puts it back into the delivered/active lists depending on status)."""
    user = request.user

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, id=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, id=container_id, company=company)

    container.is_archived = False
    container.archived_at = None
    container.save(update_fields=["is_archived", "archived_at", "updated_at"])

    messages.success(request, "Order unarchived.")
    return redirect(request.META.get("HTTP_REFERER") or "order_tracker")


@login_required
def order_tracker_archived_view(request):
    """Archived orders list (read-only tracking view)."""
    user = request.user

    if user.is_superuser:
        company = None
        containers = OrderContainer.objects.filter(is_archived=True).prefetch_related("tags")
    else:
        company = get_object_or_404(Company, owner=user)
        containers = OrderContainer.objects.filter(company=company, is_archived=True).prefetch_related("tags")

    q = (request.GET.get("q") or "").strip()
    status = (request.GET.get("status") or "").strip()
    assigned_to = (request.GET.get("assigned_to") or "").strip()
    sort = (request.GET.get("sort") or "").strip()
    direction = (request.GET.get("dir") or "desc").strip().lower()

    if q:
        containers = containers.filter(
            Q(customer_name__icontains=q)
            | Q(location_name__icontains=q)
            | Q(po_number__icontains=q)
            | Q(rpc_number__icontains=q)
            | Q(container_number__icontains=q)
            | Q(carrier__icontains=q)
            | Q(booking_number__icontains=q)
            | Q(bill_of_lading_number__icontains=q)
        )
    if status:
        containers = containers.filter(status__icontains=status)
    if assigned_to:
        containers = containers.filter(assigned_to__iexact=assigned_to)

    sort_map = {
        "customer": "customer_name",
        "location": "location_name",
        "requested": "requested_date",
        "eta": "eta",
        "est_delivery": "estimated_delivery_date",
        "loading": "loading_date",
        "assigned_to": "assigned_to",
    }
    sort_field = sort_map.get(sort)
    if direction not in {"asc", "desc"}:
        direction = "desc"

    if sort_field:
        prefix = "" if direction == "asc" else "-"
        containers = containers.order_by(f"{prefix}{sort_field}", "-updated_at", "-created_at")
    else:
        containers = containers.order_by("-updated_at", "-created_at")

    return render(
        request,
        "core/order_tracker_archived.html",
        {
            "automation_name": "Order Tracker",
            "company": company,
            "company_display": (company.name if company else "All Companies"),
            "containers": containers,
            "q": q,
            "status": status,
            "assigned_to": assigned_to,
            "sort": sort,
            "dir": direction,
        },
    )


def _ordinal(n: int) -> str:
    """Return an ordinal suffix for a positive day number (1 -> '1st')."""
    if 10 <= (n % 100) <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suffix}"


def _fmt_long_date(d: dt.date | None) -> str:
    """Format like 'January 16th 2026' or 'TBD'."""
    if not d:
        return "TBD"
    return f"{d.strftime('%B')} {_ordinal(d.day)} {d.year}"


def _fmt_short_date(d: dt.date | None) -> str:
    """Format like '12/8/2025' or 'TBD'."""
    if not d:
        return "TBD"
    # avoid platform-specific %-m / %-d
    return f"{d.month}/{d.day}/{d.year}"


@login_required
def order_tracker_sync_jsoncargo_view(request):
    """Launch the same safer JSONCargo sync command the user runs in Render Shell.

    The previous in-request button path could produce a browser 500/timeout even
    when the Render Shell command worked. This starts the exact management-shell
    command as a separate process and immediately returns to the tracker; stdout
    and tracebacks go to Render logs.
    """
    if request.method != "POST":
        return HttpResponseForbidden("POST required")

    import subprocess
    import sys

    api_key = os.getenv("JSONCARGO_API_KEY", "").strip()
    if not api_key:
        messages.error(request, "JSONCARGO_API_KEY is not set on the server.")
        return redirect("order_tracker")

    command = """import traceback; from core.services.jsoncargo_order_tracker import sync_all_containers; print('Starting JSONCargo sync...');
try:
    result = sync_all_containers()
    print('JSONCargo sync complete:', result)
except Exception as e:
    print('JSONCargo sync failed:', type(e).__name__, str(e))
    traceback.print_exc()"""

    try:
        subprocess.Popen(
            [sys.executable, "manage.py", "shell", "-c", command],
            cwd=str(settings.BASE_DIR),
            stderr=subprocess.STDOUT,
        )
        messages.success(
            request,
            "JSONCargo sync started using the same safer Render Shell command. Refresh the tracker after it finishes; details will appear in Render logs.",
        )
    except Exception as e:
        messages.error(request, f"Could not start JSONCargo sync command: {type(e).__name__}: {e}")

    return redirect("order_tracker")


@require_POST
@login_required
def order_tracker_clear_jsoncargo_updates_view(request):
    """Clear all pending JSONCargo updates/errors for the current user's visible orders."""
    from core.models import OrderContainerTrackingUpdate

    user = request.user
    if user.is_superuser:
        containers = OrderContainer.objects.filter(is_archived=False)
    else:
        company = get_object_or_404(Company, owner=user)
        containers = OrderContainer.objects.filter(company=company, is_archived=False)

    deleted_count, _details = OrderContainerTrackingUpdate.objects.filter(
        container__in=containers,
    ).delete()

    messages.success(request, f"Cleared {deleted_count} JSONCargo update record(s).")
    return redirect(request.META.get("HTTP_REFERER") or "order_tracker")


@require_POST
@login_required
def order_container_sync_jsoncargo_view(request, container_id: int):
    """Track one specific container with JSONCargo and create/update its pending update."""
    import traceback
    from core.services.jsoncargo_order_tracker import sync_one_container

    user = request.user
    if user.is_superuser:
        container = get_object_or_404(OrderContainer, id=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, id=container_id, company=company)

    if getattr(container, "is_archived", False):
        messages.error(request, "This order is archived. Unarchive it before syncing.")
        return redirect("order_container_edit", container_id=container.id)

    api_key = os.getenv("JSONCARGO_API_KEY", "").strip()
    if not api_key:
        messages.error(request, "JSONCARGO_API_KEY is not set on the server.")
        return redirect("order_container_edit", container_id=container.id)

    print(f"Starting JSONCargo sync for container {container.container_number}...")
    try:
        result, pending = sync_one_container(container, api_key=api_key)
        print("Single-container JSONCargo sync complete:", result)
        if pending and getattr(pending, "note", ""):
            messages.success(request, f"JSONCargo check complete for {container.container_number}: {result}. {pending.note}")
        else:
            messages.success(request, f"JSONCargo check complete for {container.container_number}: {result}.")
    except Exception as e:
        print("Single-container JSONCargo sync failed:", type(e).__name__, str(e))
        traceback.print_exc()
        messages.error(request, f"JSONCargo sync failed for {container.container_number}: {type(e).__name__}: {e}")

    return redirect(request.META.get("HTTP_REFERER") or reverse("order_container_edit", kwargs={"container_id": container.id}))


@require_POST
@login_required
def order_tracker_bulk_update_view(request):
    """Save manual ETA edits made directly on the master Order Tracker screen."""
    user = request.user
    if user.is_superuser:
        containers = OrderContainer.objects.filter(is_archived=False).exclude(status__iexact="Delivered")
    else:
        company = get_object_or_404(Company, owner=user)
        containers = OrderContainer.objects.filter(company=company, is_archived=False).exclude(status__iexact="Delivered")

    ids = request.POST.getlist("container_ids")
    if ids:
        containers = containers.filter(id__in=ids)

    changed_count = 0
    error_count = 0

    for c in containers:
        raw_eta = (request.POST.get(f"eta_{c.id}") or "").strip()
        raw_city = (request.POST.get(f"eta_city_{c.id}") or "").strip()

        try:
            new_eta = datetime.datetime.strptime(raw_eta, "%Y-%m-%d").date() if raw_eta else None
        except ValueError:
            error_count += 1
            continue

        fields = []
        if c.eta != new_eta:
            c.eta = new_eta
            fields.append("eta")
        if (c.eta_city or "") != raw_city:
            c.eta_city = raw_city
            fields.append("eta_city")

        if fields:
            fields.append("updated_at")
            c.save(update_fields=fields)
            changed_count += 1

    if error_count:
        messages.error(request, f"Saved {changed_count} ETA edit(s). {error_count} row(s) had invalid date format; use YYYY-MM-DD.")
    else:
        messages.success(request, f"Saved {changed_count} ETA edit(s).")

    return redirect(request.META.get("HTTP_REFERER") or "order_tracker")

def order_tracker_recap_docx_view(request):
    """Download an Order Recap DOCX for all (filtered) containers.

    Output is ordered by RPC# (requested), with reasonable tie-breakers.
    """
    user = request.user

    # Superusers see ALL containers across companies.
    if user.is_superuser:
        containers = OrderContainer.objects.filter(is_archived=False)
    else:
        company = get_object_or_404(Company, owner=user)
        containers = OrderContainer.objects.filter(company=company, is_archived=False)

    # Respect the same filters as the dashboard (if present in the URL).
    q = (request.GET.get("q") or "").strip()
    status = (request.GET.get("status") or "").strip()
    assigned_to = (request.GET.get("assigned_to") or "").strip()

    if q:
        containers = containers.filter(
            Q(customer_name__icontains=q)
            | Q(location_name__icontains=q)
            | Q(po_number__icontains=q)
            | Q(rpc_number__icontains=q)
            | Q(booking_number__icontains=q)
            | Q(bill_of_lading_number__icontains=q)
            | Q(container_number__icontains=q)
            | Q(carrier__icontains=q)
        )
    if status:
        containers = containers.filter(status__icontains=status)
    if assigned_to:
        containers = containers.filter(assigned_to__iexact=assigned_to)

    # Prefetch lines first; we'll do a Python-side sort so we can sort by the
    # leading numeric portion of RPC# even when rpc_number includes extra text
    # like "6066 Miami".
    containers = containers.prefetch_related("lines")

    def _rpc_sort_key(raw: str | None) -> tuple[int, str]:
        """Return (numeric_rpc, raw_rpc) for sorting.

        Examples:
        - "6066 Miami" -> (6066, "6066 Miami")
        - "6066" -> (6066, "6066")
        - "TBD"/None/"" -> (10**9, "")  (push to end)
        """
        s = (raw or "").strip()
        if not s:
            return (10**9, "")
        first = s.split()[0]
        digits = "".join(ch for ch in first if ch.isdigit())
        if digits:
            try:
                return (int(digits), s)
            except Exception:
                pass
        return (10**9, s)

    # Sort by Customer, then Location, then RPC# (numeric), with tie-breakers for stability.
    containers = sorted(
        list(containers),
        key=lambda c: (
            (getattr(c, "customer_name", "") or "").strip().lower(),
            (getattr(c, "location_name", "") or "").strip().lower(),
            _rpc_sort_key(getattr(c, "rpc_number", None)),
            getattr(c, "requested_date", None) or dt.date.min,
            (getattr(c, "po_number", "") or "").strip(),
            -(getattr(c, "updated_at", None).timestamp() if getattr(c, "updated_at", None) else 0),
        ),
    )

    doc = Document()

    # Color palette + formatting to match the recap example document.
    # (Word RGB hex values)
    COLOR_RED = RGBColor(0xC0, 0x00, 0x00)
    COLOR_BROWN = RGBColor(0x7F, 0x60, 0x00)
    COLOR_ORANGE = RGBColor(0xE6, 0x91, 0x38)
    COLOR_BLUE = RGBColor(0x05, 0x63, 0xC1)
    COLOR_GREEN = RGBColor(0x00, 0xB0, 0x50)
    COLOR_BLACK = RGBColor(0x11, 0x11, 0x11)

    def _add_run(p, text: str, color: RGBColor, *, bold: bool | None = None, underline: bool | None = None, highlight=None):
        r = p.add_run(text)
        r.font.color.rgb = color
        if bold is not None:
            r.bold = bold
        if underline is not None:
            r.font.underline = underline
        if highlight is not None:
            r.font.highlight_color = highlight
        return r

    # Match a clean, simple default look.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for idx, c in enumerate(containers):
        po = (c.po_number or "TBD").strip() or "TBD"
        rpc = (c.rpc_number or "TBD").strip() or "TBD"
        status_txt = (c.status or "TBD").strip() or "TBD"

        # Line 1: PO + Requested + Status (with colors)
        p1 = doc.add_paragraph()
        _add_run(p1, f"PO#{po}– ", COLOR_RED, bold=True)
        _add_run(p1, f"Requested {_fmt_long_date(c.requested_date)} ", COLOR_BROWN)
        # STATUS: orange + bold + highlighted (user requested highlight)
        _add_run(p1, f"***{status_txt}*** ", COLOR_ORANGE, bold=True, highlight=WD_COLOR_INDEX.YELLOW)

        # Line 2: RPC (blue, bold)
        p2 = doc.add_paragraph()
        # IMPORTANT: the RPC string already includes the city (ex: "6066 Miami"),
        # so do NOT append the location name.
        _add_run(p2, f"RPC# {rpc} ", COLOR_BLUE, bold=True)

        # Line 3: Loading Date (green, bold)
        p3 = doc.add_paragraph()
        _add_run(p3, f"Loading Date: {_fmt_short_date(c.loading_date)}", COLOR_GREEN, bold=True)

        # Content lines (green, underlined like example)
        # NOTE: user requested to remove the literal "Buckets/Pallets" label.
        for line in c.lines.all():
            desc = (line.item_description or "").strip() or "(item)"
            pallets = int(line.pallets or 0)
            upp = int(line.units_per_pallet or 0)
            total = int(line.total_units or (pallets * upp))

            p = doc.add_paragraph()
            # Description part (underlined in the example as well)
            _add_run(p, f"{desc}  (", COLOR_GREEN, underline=True)
            _add_run(p, f"{pallets} x {upp} = {total:,} ", COLOR_GREEN, underline=True)
            _add_run(p, "Pieces", COLOR_GREEN, underline=True)
            _add_run(p, ")", COLOR_GREEN, underline=True)

        # Dates (black)
        p = doc.add_paragraph()
        _add_run(p, f"ETD: {_fmt_short_date(c.etd)}", COLOR_BLACK)
        p = doc.add_paragraph()
        _add_run(p, f"ETA: {_fmt_short_date(c.eta)}", COLOR_BLACK)
        p = doc.add_paragraph()
        _add_run(p, f"Estimated Delivery: {_fmt_short_date(c.estimated_delivery_date)}", COLOR_BLACK)

        # spacing between containers
        if idx != len(containers) - 1:
            doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"order_recap_{timezone.localdate().isoformat()}.docx"
    return FileResponse(
        buf,
        as_attachment=True,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
@require_http_methods(["GET", "POST"])
def order_container_commercial_invoice_view(request, container_id: int):
    """Build a commercial invoice for one order container.

    The invoice uses the order's PO#, loading date, and packing lines. The user
    enters the bill-to address, Euro pricing for each line, and chooses the
    commercial terms blurb before downloading a DOCX invoice.
    """
    user = request.user
    if user.is_superuser:
        container = get_object_or_404(OrderContainer.objects.prefetch_related("lines"), pk=container_id)
        company = container.company
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(
            OrderContainer.objects.prefetch_related("lines"),
            pk=container_id,
            company=company,
        )

    lines = list(container.lines.all())

    def _attach_invoice_price_values(post_data=None):
        post_data = post_data or {}
        for ln in lines:
            try:
                ln.invoice_price = post_data.get(f"price_{ln.id}", "")
            except Exception:
                ln.invoice_price = ""

    _attach_invoice_price_values()

    if not lines:
        messages.error(request, "Add at least one packing line before creating a commercial invoice.")
        return redirect("order_container_edit", container_id=container.id)

    BLURB_DELIVERED = "Delivered prices. The sales conditions of Retriever Packaging Company LLC are applicable."
    BLURB_EXW = "EXW Waalwijk. Sales conditions of Retriever Packaging Company LLC are applicable."
    blurb_choices = [
        ("delivered", BLURB_DELIVERED),
        ("exw", BLURB_EXW),
        ("none", "No bottom blurb"),
    ]

    if request.method == "POST":
        bill_to = (request.POST.get("bill_to") or "").strip()
        invoice_number = (request.POST.get("invoice_number") or "").strip()
        selected_blurb_key = (request.POST.get("terms_blurb") or "delivered").strip()
        if selected_blurb_key == "none":
            terms_blurb = ""
        elif selected_blurb_key == "exw":
            terms_blurb = BLURB_EXW
        else:
            terms_blurb = BLURB_DELIVERED

        if not bill_to:
            _attach_invoice_price_values(request.POST)
            messages.error(request, "Please enter a bill-to address.")
            return render(
                request,
                "core/commercial_invoice_form.html",
                {
                    "automation_name": "Commercial Invoice",
                    "company": company,
                    "container": container,
                    "lines": lines,
                    "blurb_choices": blurb_choices,
                    "selected_blurb": selected_blurb_key,
                    "bill_to": bill_to,
                    "invoice_number": invoice_number,
                    "posted_prices": request.POST,
                },
            )

        priced_lines = []
        errors = []
        grand_total = Decimal("0.00")
        for line in lines:
            raw_price = (request.POST.get(f"price_{line.id}") or "").strip().replace(",", ".")
            try:
                unit_price = Decimal(raw_price).quantize(Decimal("0.0001"))
                if unit_price < 0:
                    raise InvalidOperation()
            except Exception:
                errors.append(f"Enter a valid Euro price for {line.item_description}.")
                unit_price = Decimal("0.0000")

            pallets = int(line.pallets or 0)
            units_per_pallet = int(line.units_per_pallet or 0)
            total_units = int(line.total_units or (pallets * units_per_pallet))
            amount = (Decimal(total_units) * unit_price).quantize(Decimal("0.01"))
            grand_total += amount
            priced_lines.append(
                {
                    "line": line,
                    "pallets": pallets,
                    "units_per_pallet": units_per_pallet,
                    "total_units": total_units,
                    "unit_price": unit_price,
                    "amount": amount,
                }
            )

        if errors:
            _attach_invoice_price_values(request.POST)
            for err in errors:
                messages.error(request, err)
            return render(
                request,
                "core/commercial_invoice_form.html",
                {
                    "automation_name": "Commercial Invoice",
                    "company": company,
                    "container": container,
                    "lines": lines,
                    "blurb_choices": blurb_choices,
                    "selected_blurb": selected_blurb_key,
                    "bill_to": bill_to,
                    "invoice_number": invoice_number,
                    "posted_prices": request.POST,
                },
            )

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        def add_label_value(paragraph, label: str, value: str):
            r = paragraph.add_run(label)
            r.bold = True
            paragraph.add_run(value)

        title = doc.add_paragraph()
        title.alignment = 1  # centered
        run = title.add_run("COMMERCIAL INVOICE")
        run.bold = True
        run.font.size = Pt(16)

        header = doc.add_table(rows=1, cols=2)
        header.style = "Table Grid"
        left = header.cell(0, 0)
        right = header.cell(0, 1)
        p_left = left.paragraphs[0]
        add_label_value(p_left, "Shipper:\n", "Naber Plastics BV\nVan Hilststraat 12\n5145 RL Waalwijk\nThe Netherlands")
        p_right = right.paragraphs[0]
        invoice_date = container.loading_date or timezone.localdate()
        add_label_value(p_right, "Invoice date: ", _fmt_short_date(invoice_date))
        p = right.add_paragraph()
        add_label_value(p, "PO#: ", (container.po_number or "TBD"))
        p = right.add_paragraph()
        add_label_value(p, "RPC#: ", (container.rpc_number or "TBD"))
        if invoice_number:
            p = right.add_paragraph()
            add_label_value(p, "Invoice #: ", invoice_number)
        if container.container_number:
            p = right.add_paragraph()
            add_label_value(p, "Container #: ", container.container_number)

        doc.add_paragraph("")
        addr = doc.add_table(rows=1, cols=2)
        addr.style = "Table Grid"
        bill_cell = addr.cell(0, 0)
        ship_cell = addr.cell(0, 1)
        add_label_value(bill_cell.paragraphs[0], "Bill To:\n", bill_to)
        ship_to_parts = [container.customer_name]
        if container.location_name:
            ship_to_parts.append(container.location_name)
        add_label_value(ship_cell.paragraphs[0], "Order / Destination:\n", "\n".join(ship_to_parts))

        p = doc.add_paragraph()
        add_label_value(p, "Ship via: ", "sea container")
        p = doc.add_paragraph()
        add_label_value(p, "Loading date: ", _fmt_short_date(container.loading_date))

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        headings = ["Description", "Packing", "Pallets", "Quantity", "Unit price EUR", "Amount EUR"]
        for idx, heading in enumerate(headings):
            hdr[idx].text = heading
            for run in hdr[idx].paragraphs[0].runs:
                run.bold = True

        for item in priced_lines:
            line = item["line"]
            row = table.add_row().cells
            row[0].text = line.item_description or ""
            row[1].text = f'{item["pallets"]} x {item["units_per_pallet"]} pieces/pallet'
            row[2].text = f'{item["pallets"]:,}'
            row[3].text = f'{item["total_units"]:,}'
            row[4].text = f'€ {item["unit_price"]:,.4f}'
            row[5].text = f'€ {item["amount"]:,.2f}'

        total_row = table.add_row().cells
        total_row[0].text = "Total"
        total_row[1].text = ""
        total_row[2].text = ""
        total_row[3].text = f'{sum(i["total_units"] for i in priced_lines):,}'
        total_row[4].text = ""
        total_row[5].text = f'€ {grand_total:,.2f}'
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        if terms_blurb:
            doc.add_paragraph("")
            p = doc.add_paragraph(terms_blurb)
            p.runs[0].bold = True

        doc.add_paragraph("Country of origin: The Netherlands")
        doc.add_paragraph("Currency: EUR")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe_po = re.sub(r"[^A-Za-z0-9_-]+", "_", (container.po_number or f"order_{container.id}")).strip("_")
        filename = f"commercial_invoice_{safe_po}.docx"
        return FileResponse(
            buf,
            as_attachment=True,
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    return render(
        request,
        "core/commercial_invoice_form.html",
        {
            "automation_name": "Commercial Invoice",
            "company": company,
            "container": container,
            "lines": lines,
            "blurb_choices": blurb_choices,
            "selected_blurb": "delivered",
            "bill_to": "",
            "invoice_number": "",
            "posted_prices": {},
        },
    )


@login_required
@require_http_methods(["POST"])
def order_container_toggle_delivered_view(request, container_id: int):
    """Toggle an OrderContainer between Delivered and Active."""
    user = request.user

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, pk=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, pk=container_id, company=company)

    current = (container.status or "").strip()
    if current.lower() == "delivered":
        # Re-open as active
        container.status = "In transit"
        messages.success(request, "Order marked active.")
    else:
        container.status = "Delivered"
        messages.success(request, "Order marked delivered.")

    container.save(update_fields=["status", "updated_at"])
    return redirect("order_tracker")


@login_required
def order_container_tracking_approve_view(request, container_id: int, update_id: int):
    """Approve a pending JSONCargo tracking update and apply it to the container."""
    if request.method != "POST":
        return HttpResponseForbidden("POST required")

    user = request.user
    from core.models import OrderContainerTrackingUpdate

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, pk=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, pk=container_id, company=company)

    upd = get_object_or_404(
        OrderContainerTrackingUpdate,
        pk=update_id,
        container=container,
        status=OrderContainerTrackingUpdate.STATUS_PENDING,
    )

    if getattr(container, "is_archived", False):
        messages.error(request, "This order is archived. Unarchive it to apply tracking updates.")
        return redirect("order_container_edit", container_id=container.id)

    # Approve behavior depends on kind.
    if getattr(upd, "kind", OrderContainerTrackingUpdate.KIND_CHANGE) in (
        OrderContainerTrackingUpdate.KIND_NO_CHANGE,
        getattr(OrderContainerTrackingUpdate, "KIND_ERROR", "error"),
    ):
        # Require explicit acknowledgement checkbox
        if request.POST.get("acknowledge_no_change") != "1":
            messages.error(request, "Please check the acknowledgement box before approving.")
            return redirect("order_container_edit", container_id=container.id)

        # No fields are changed; we just mark the note as approved.
        upd.status = OrderContainerTrackingUpdate.STATUS_APPROVED
        upd.decided_by = user
        upd.decided_at = timezone.now()
        upd.save(update_fields=["status", "decided_by", "decided_at", "updated_at"])

        if getattr(upd, "kind", "") == getattr(OrderContainerTrackingUpdate, "KIND_ERROR", "error"):
            messages.success(request, "Tracking error acknowledged.")
        else:
            messages.success(request, "API check acknowledged (no change from current tracker).")
        return redirect("order_container_edit", container_id=container.id)

    # KIND_CHANGE: Apply proposed values
    changed = False
    if upd.proposed_eta and upd.proposed_eta != container.eta:
        container.eta = upd.proposed_eta
        changed = True
    if (upd.proposed_eta_city or "") != (container.eta_city or ""):
        container.eta_city = upd.proposed_eta_city or ""
        changed = True

    if changed:
        container.save(update_fields=["eta", "eta_city", "updated_at"])

    upd.status = OrderContainerTrackingUpdate.STATUS_APPROVED
    upd.decided_by = user
    upd.decided_at = timezone.now()
    upd.save(update_fields=["status", "decided_by", "decided_at", "updated_at"])

    messages.success(request, "Tracking update approved and applied.")
    return redirect("order_container_edit", container_id=container.id)


@login_required
def order_container_tracking_reject_view(request, container_id: int, update_id: int):
    """Reject a pending JSONCargo tracking update."""
    if request.method != "POST":
        return HttpResponseForbidden("POST required")

    user = request.user
    from core.models import OrderContainerTrackingUpdate

    if user.is_superuser:
        container = get_object_or_404(OrderContainer, pk=container_id)
    else:
        company = get_object_or_404(Company, owner=user)
        container = get_object_or_404(OrderContainer, pk=container_id, company=company)

    upd = get_object_or_404(
        OrderContainerTrackingUpdate,
        pk=update_id,
        container=container,
        status=OrderContainerTrackingUpdate.STATUS_PENDING,
    )

    if getattr(container, "is_archived", False):
        messages.error(request, "This order is archived. Unarchive it to manage tracking updates.")
        return redirect("order_container_edit", container_id=container.id)

    upd.status = OrderContainerTrackingUpdate.STATUS_REJECTED
    upd.decided_by = user
    upd.decided_at = timezone.now()
    upd.save(update_fields=["status", "decided_by", "decided_at", "updated_at"])

    messages.success(request, "Tracking update rejected.")
    return redirect("order_container_edit", container_id=container.id)


@login_required
def order_container_edit_view(request, container_id: int | None = None):
    """
    Create or edit a container + its content lines (inline formset).
    """
    user = request.user

    # IMPORTANT FIX:
    # Superusers can edit any container, so infer company from the container when editing.
    if user.is_superuser:
        if container_id is None:
            company = Company.objects.order_by("id").first()
            container = None
        else:
            container = get_object_or_404(OrderContainer, pk=container_id)
            company = container.company
    else:
        company = get_object_or_404(Company, owner=user)
        if container_id is None:
            container = None
        else:
            container = get_object_or_404(OrderContainer, pk=container_id, company=company)

    LineFormSet = inlineformset_factory(
        parent_model=OrderContainer,
        model=OrderContainerLine,
        form=OrderContainerLineForm,
        extra=3,
        can_delete=True,
    )

    DocumentFormSet = inlineformset_factory(
        parent_model=OrderContainer,
        model=OrderContainerDocument,
        form=OrderContainerDocumentForm,
        extra=1,
        can_delete=True,
    )

    TagFormSet = inlineformset_factory(
        parent_model=OrderContainer,
        model=OrderContainerTag,
        form=OrderContainerTagForm,
        extra=2,
        can_delete=True,
    )

    # If this endpoint is hit by the "Quick add" form on the dashboard, the POST
    # will *not* include the formset management fields (lines-TOTAL_FORMS, etc.).
    # In that case we should create the container using just OrderContainerForm
    # and then redirect into the full editor.
    def _is_quick_add_post(req):
        if req.method != "POST":
            return False
        return (
            "lines-TOTAL_FORMS" not in req.POST
            and "docs-TOTAL_FORMS" not in req.POST
            and "tags-TOTAL_FORMS" not in req.POST
        )

    if request.method == "POST":
        if _is_quick_add_post(request):
            # Quick add: create container without formsets, then redirect to full editor.
            form = OrderContainerForm(request.POST, instance=container)
            if form.is_valid():
                obj: OrderContainer = form.save(commit=False)
                obj.company = company
                if obj.created_by_id is None:
                    obj.created_by = user
                obj.save()
                messages.success(request, "Container created. Add lines/documents below.")
                return redirect("order_container_edit", container_id=obj.id)
            # If the minimal form is invalid, fall through and render with errors.
            formset = LineFormSet(instance=container, prefix="lines")
            doc_formset = DocumentFormSet(instance=container, prefix="docs")
            tag_formset = TagFormSet(instance=container, prefix="tags")
            messages.error(request, "Please fix the errors below and try again.")
        else:
            form = OrderContainerForm(request.POST, instance=container)
            formset = LineFormSet(request.POST, instance=container, prefix="lines")
            doc_formset = DocumentFormSet(
                request.POST,
                request.FILES,
                instance=container,
                prefix="docs",
            )
            tag_formset = TagFormSet(request.POST, instance=container, prefix="tags")
            if form.is_valid() and formset.is_valid() and doc_formset.is_valid() and tag_formset.is_valid():
                obj: OrderContainer = form.save(commit=False)
                obj.company = company
                if obj.created_by_id is None:
                    obj.created_by = user
                obj.save()

                formset.instance = obj
                formset.save()

                doc_formset.instance = obj
                doc_formset.save()

                tag_formset.instance = obj
                tag_formset.save()

                # Optional: user acknowledged the latest pending JSONCargo update while saving.
                ack_id = (request.POST.get("ack_tracking_update_id") or "").strip()
                if ack_id and container is not None:
                    try:
                        from core.models import OrderContainerTrackingUpdate
                        pending = OrderContainerTrackingUpdate.objects.filter(
                            id=int(ack_id),
                            container=obj,
                            status=OrderContainerTrackingUpdate.STATUS_PENDING,
                        ).first()
                        if pending:
                            pending.status = OrderContainerTrackingUpdate.STATUS_APPROVED
                            pending.decided_by = user
                            pending.decided_at = timezone.now()
                            pending.note = (pending.note or "").strip()
                            pending.save(update_fields=[
                                "status",
                                "decided_by",
                                "decided_at",
                                "note",
                                "updated_at",
                            ])
                    except Exception:
                        # Don't block saving the container if the ack fails for any reason.
                        pass

                messages.success(request, "Container saved.")
                return redirect("order_container_edit", container_id=obj.id)
            else:
                messages.error(request, "Please fix the errors below and try again.")
    else:
        form = OrderContainerForm(instance=container)
        formset = LineFormSet(instance=container, prefix="lines")
        doc_formset = DocumentFormSet(instance=container, prefix="docs")
        tag_formset = TagFormSet(instance=container, prefix="tags")

    # Latest pending JSONCargo update (requires approval)
    pending_tracking_update = None
    if container is not None and not getattr(container, "is_archived", False):
        try:
            from core.models import OrderContainerTrackingUpdate

            pending_tracking_update = (
                OrderContainerTrackingUpdate.objects.filter(
                    container=container,
                    status=OrderContainerTrackingUpdate.STATUS_PENDING,
                )
                .order_by("-created_at", "-id")
                .first()
            )
        except Exception:
            pending_tracking_update = None

    
    # JSONCargo snapshot (flat "data" object)
    # We only expose a small, human-friendly subset in the template for manual review.
    jsoncargo_data: dict = {}
    jsoncargo_next_destination: str = ""
    jsoncargo_display_eta: str = ""
    if pending_tracking_update is not None:
        payload = getattr(pending_tracking_update, "source_payload", None) or {}
        if isinstance(payload, dict):
            raw_data = payload.get("data")
            if isinstance(raw_data, dict):
                jsoncargo_data = raw_data

    # Best-effort "next destination" label (city/port). JSONCargo may leave next_location null;
    # fall back to other routing fields so the UI always shows something useful.
    def _clean_str(val) -> str:
        if val is None:
            return ""
        return str(val).strip()

    if isinstance(jsoncargo_data, dict):
        for key in (
            "eta_next_destination",
            "eta_final_destination",
            "eta_destination",
            "eta",
            "eta_delivery",
            "eta_discharge",
        ):
            candidate = _clean_str(jsoncargo_data.get(key))
            if candidate:
                jsoncargo_display_eta = candidate
                break

        for key in (
            "next_location",
            "next_location_terminal",
            "final_destination",
            "final_destination_port",
            "final_destination_city",
            "destination",
            "delivery_to",
            "delivered_to",
            "consignee_city",
            "shipped_to",
            "shipped_to_terminal",
            "discharging_port",
        ):
            candidate = _clean_str(jsoncargo_data.get(key))
            if candidate:
                jsoncargo_next_destination = candidate
                break

    return render(
        request,
        "core/order_container_edit.html",
        {
            "automation_name": "Order Tracker",
            "company": company,
            "container": container,
            "form": form,
            "formset": formset,
            "doc_formset": doc_formset,
            "tag_formset": tag_formset,
            "pending_tracking_update": pending_tracking_update,
            "jsoncargo_data": jsoncargo_data,
            "jsoncargo_display_eta": jsoncargo_display_eta,
            "jsoncargo_next_destination": jsoncargo_next_destination,
        },
    )















# =========================
# Schedule Dashboard
# =========================

@login_required
def schedule_dashboard_view(request):
    """Company scheduling dashboard (week view).

    GET params:
      - d: reference date (YYYY-MM-DD). The week shown is the Monday-starting week containing d.
      - view: "week" (default) or "month"
    """
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    d_str = (request.GET.get("d") or "").strip()
    try:
        ref_date = dt.date.fromisoformat(d_str) if d_str else timezone.localdate()
    except Exception:
        ref_date = timezone.localdate()

    view_mode = (request.GET.get("view") or "week").strip().lower()
    view_mode = "month" if view_mode == "month" else "week"

    # Always-on notes (not tied to a date)
    global_note_obj, _ = ScheduleGlobalNote.objects.get_or_create(company=company)
    global_note_form = ScheduleGlobalNoteForm(instance=global_note_obj)

    form = ScheduleActivityForm(initial={"date": ref_date, "repeat_every": 1, "repeat_unit": "weeks"})

    def _expand_occurrences(qs, start_date: dt.date, end_date: dt.date):
        """Return a list of *display* activities for the given date range.

        For recurring activities, we generate per-occurrence copies so that the
        same base object can appear on multiple days without overwriting state.
        """
        out = []
        for a in qs:
            if not a.is_recurring:
                out.append(copy.copy(a))
                continue

            until = a.repeat_until or end_date
            occ_end = min(end_date, until)
            if a.date > occ_end:
                continue

            unit = (a.repeat_unit or "weeks").lower()
            every = int(a.repeat_every or 1)
            if every < 1:
                every = 1

            cur = a.date

            # Advance to first occurrence on/after start_date
            if unit == "days":
                step = dt.timedelta(days=every)
                while cur < start_date:
                    cur += step
            elif unit == "weeks":
                step = dt.timedelta(days=7 * every)
                while cur < start_date:
                    cur += step
            else:  # months
                while cur < start_date:
                    cur = cur + relativedelta(months=every)

            while cur <= occ_end:
                occ = copy.copy(a)
                occ.date = cur
                occ._is_occurrence = True  # type: ignore[attr-defined]
                out.append(occ)
                if unit == "days":
                    cur += dt.timedelta(days=every)
                elif unit == "weeks":
                    cur += dt.timedelta(days=7 * every)
                else:
                    cur = cur + relativedelta(months=every)

        return out

    if view_mode == "month":
        # Month grid (Monday-starting weeks)
        month_start = ref_date.replace(day=1)
        # Find the first Monday to show
        grid_start = month_start - dt.timedelta(days=month_start.weekday())
        # Find the last day of the month
        if month_start.month == 12:
            next_month = month_start.replace(year=month_start.year + 1, month=1, day=1)
        else:
            next_month = month_start.replace(month=month_start.month + 1, day=1)
        month_end = next_month - dt.timedelta(days=1)
        # Extend to the end of the last week (Sunday)
        grid_end = month_end + dt.timedelta(days=(6 - month_end.weekday()))

        base_qs = (
            ScheduleActivity.objects
            .filter(company=company)
            .filter(
                Q(is_recurring=False, date__gte=grid_start, date__lte=grid_end)
                | (
                    Q(is_recurring=True, date__lte=grid_end)
                    & (Q(repeat_until__isnull=True) | Q(repeat_until__gte=grid_start))
                )
            )
            .order_by("date", "start_time", "created_at", "id")
        )

        display_acts = _expand_occurrences(base_qs, grid_start, grid_end)
        display_acts.sort(key=lambda a: (a.date, a.start_time or dt.time(23, 59), a.created_at, a.id))

        by_day = {}
        for a in display_acts:
            by_day.setdefault(a.date, []).append(a)

        weeks = []
        cur = grid_start
        while cur <= grid_end:
            week_days = [cur + dt.timedelta(days=i) for i in range(7)]
            weeks.append(
                [
                    {
                        "date": d,
                        "in_month": (d.month == month_start.month),
                        "activities": by_day.get(d, []),
                    }
                    for d in week_days
                ]
            )
            cur += dt.timedelta(days=7)

        # Month nav
        prev_month = (month_start - dt.timedelta(days=1)).replace(day=1)
        next_month = (month_end + dt.timedelta(days=1)).replace(day=1)

        context = {
            "automation_name": "Schedule Dashboard",
            "company": company,
            "ref_date": ref_date,
            "view_mode": "month",
            "month_start": month_start,
            "grid_start": grid_start,
            "grid_end": grid_end,
            "prev_d": prev_month.isoformat(),
            "next_d": next_month.isoformat(),
            "weeks": weeks,
            "form": form,
            "global_note_form": global_note_form,
            "global_note_obj": global_note_obj,
        }
        return render(request, "core/schedule_month.html", context)

    # Week view (default)
    week_start = ref_date - dt.timedelta(days=ref_date.weekday())
    days = [week_start + dt.timedelta(days=i) for i in range(7)]
    week_end = days[-1]

    base_qs = (
        ScheduleActivity.objects
        .filter(company=company)
        .filter(
            Q(is_recurring=False, date__gte=week_start, date__lte=week_end)
            | (
                Q(is_recurring=True, date__lte=week_end)
                & (Q(repeat_until__isnull=True) | Q(repeat_until__gte=week_start))
            )
        )
        .order_by("date", "start_time", "created_at", "id")
    )

    display_acts = _expand_occurrences(base_qs, week_start, week_end)
    display_acts.sort(key=lambda a: (a.date, a.start_time or dt.time(23, 59), a.created_at, a.id))

    by_day = {d: [] for d in days}
    for a in display_acts:
        by_day.setdefault(a.date, []).append(a)

    day_blocks = [{"date": d, "activities": by_day.get(d, [])} for d in days]

    context = {
        "automation_name": "Schedule Dashboard",
        "company": company,
        "ref_date": ref_date,
        "view_mode": "week",
        "week_start": week_start,
        "week_end": week_end,
        "prev_d": (week_start - dt.timedelta(days=7)).isoformat(),
        "next_d": (week_start + dt.timedelta(days=7)).isoformat(),
        "day_blocks": day_blocks,
        "form": form,
        "global_note_form": global_note_form,
        "global_note_obj": global_note_obj,
    }
    return render(request, "core/schedule_dashboard.html", context)


@login_required
@require_POST
def schedule_activity_add_view(request):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    form = ScheduleActivityForm(request.POST)
    if form.is_valid():
        obj = form.save(commit=False)
        obj.company = company
        obj.save()
        messages.success(request, "Activity added.")
    else:
        messages.error(request, "Could not add activity. Please check the fields.")

    back_d = (request.POST.get("back_d") or "").strip()
    back_view = (request.POST.get("back_view") or "").strip().lower()
    back_view = "month" if back_view == "month" else "week"
    if back_d:
        return redirect(f"/automations/schedule/?d={back_d}&view={back_view}")
    return redirect("schedule_dashboard")


@login_required
@require_POST
def schedule_global_note_save_view(request):
    """Save always-on notes for the scheduling dashboard."""
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    note_obj, _ = ScheduleGlobalNote.objects.get_or_create(company=company)
    form = ScheduleGlobalNoteForm(request.POST, instance=note_obj)
    if form.is_valid():
        form.save()
        messages.success(request, "Notes saved.")
    else:
        messages.error(request, "Could not save notes.")

    back_d = (request.POST.get("back_d") or "").strip()
    back_view = (request.POST.get("back_view") or "").strip().lower()
    back_view = "month" if back_view == "month" else "week"
    if back_d:
        return redirect(f"/automations/schedule/?d={back_d}&view={back_view}")
    return redirect("schedule_dashboard")


@login_required
def schedule_activity_edit_view(request, pk):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    activity = get_object_or_404(ScheduleActivity, pk=pk, company=company)

    if request.method == "POST":
        form = ScheduleActivityForm(request.POST, instance=activity)
        if form.is_valid():
            form.save()
            messages.success(request, "Activity updated.")
            back_d = (request.POST.get("back_d") or "").strip()
            back_view = (request.POST.get("back_view") or "").strip().lower()
            back_view = "month" if back_view == "month" else "week"
            if back_d:
                return redirect(f"/automations/schedule/?d={back_d}&view={back_view}")
            return redirect("schedule_dashboard")
    else:
        form = ScheduleActivityForm(instance=activity)

    back_d = (request.GET.get("back_d") or activity.date.isoformat()).strip()
    back_view = (request.GET.get("back_view") or "week").strip().lower()
    back_view = "month" if back_view == "month" else "week"
    return render(
        request,
        "core/schedule_activity_edit.html",
        {
            "automation_name": "Edit Activity",
            "company": company,
            "activity": activity,
            "form": form,
            "back_d": back_d,
            "back_view": back_view,
        },
    )


@login_required
@require_POST
def schedule_activity_delete_view(request, pk):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    activity = get_object_or_404(ScheduleActivity, pk=pk, company=company)
    activity.delete()
    messages.success(request, "Activity deleted.")

    back_d = (request.POST.get("back_d") or "").strip()
    back_view = (request.POST.get("back_view") or "").strip().lower()
    back_view = "month" if back_view == "month" else "week"
    if back_d:
        return redirect(f"/automations/schedule/?d={back_d}&view={back_view}")
    return redirect("schedule_dashboard")


@login_required
@require_POST
def schedule_activity_toggle_done_view(request, pk):
    user = request.user

    if user.is_superuser:
        company = Company.objects.order_by("id").first()
    else:
        company = Company.objects.filter(owner=user).order_by("id").first()

    if not company:
        return HttpResponseForbidden("No company is associated with this user.")

    activity = get_object_or_404(ScheduleActivity, pk=pk, company=company)

    if activity.status == ScheduleActivity.STATUS_DONE:
        activity.status = ScheduleActivity.STATUS_PLANNED
    else:
        activity.status = ScheduleActivity.STATUS_DONE
    activity.save(update_fields=["status", "updated_at"])

    back_d = (request.POST.get("back_d") or "").strip()
    back_view = (request.POST.get("back_view") or "").strip().lower()
    back_view = "month" if back_view == "month" else "week"
    if back_d:
        return redirect(f"/automations/schedule/?d={back_d}&view={back_view}")
    return redirect("schedule_dashboard")




























































































































































































































































































































































































































































































































































































































































































































































